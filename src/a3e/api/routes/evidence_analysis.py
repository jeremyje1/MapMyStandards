"""
Evidence analysis API routes for document processing and mapping
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks, Body
from fastapi.responses import JSONResponse
from typing import List, Optional, Dict, Any
import asyncio
import logging
import uuid
from datetime import datetime
import json
import os
from pathlib import Path

from ...services.database_service import DatabaseService
from ...services.document_service import DocumentService
from ...services.evidence_mapper import EvidenceMapper, EvidenceDocument
from ...services.evidence_processor import EvidenceProcessor
from ...services.ai_service import AIService
from ...core.config import settings
from ..dependencies import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/evidence-analysis", tags=["evidence_analysis"])

# Service instances
evidence_mapper = None
evidence_processor = None

def get_evidence_mapper() -> EvidenceMapper:
    """Get or create evidence mapper instance"""
    global evidence_mapper
    if evidence_mapper is None:
        evidence_mapper = EvidenceMapper()
    return evidence_mapper

async def get_evidence_processor() -> EvidenceProcessor:
    """Get or create evidence processor instance"""
    global evidence_processor
    if evidence_processor is None:
        ai_service = AIService(settings)
        evidence_processor = EvidenceProcessor(ai_service)
    return evidence_processor

# Database dependency
async def get_db():
    """Database dependency function"""
    db_service = DatabaseService(settings.database_url)
    async with db_service.get_session() as session:
        yield session

@router.post("/upload")
async def upload_for_analysis(
    files: List[UploadFile] = File(...),
    institution_id: Optional[str] = Form(None),
    accreditor: Optional[str] = Form(None),
    auto_analyze: bool = Form(True),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: Dict = Depends(get_current_user)
):
    """Upload documents for evidence analysis with mapping"""
    try:
        uploaded_files = []
        
        # Create temp directory for uploads
        upload_dir = Path("/tmp/evidence_uploads")
        upload_dir.mkdir(exist_ok=True)
        
        for file in files:
            # Save file temporarily
            file_id = str(uuid.uuid4())
            file_path = upload_dir / f"{file_id}_{file.filename}"
            
            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)
            
            file_info = {
                "id": file_id,
                "filename": file.filename,
                "size": len(content),
                "status": "uploaded",
                "path": str(file_path),
                "uploaded_at": datetime.utcnow().isoformat(),
                "user_id": current_user.get("user_id"),
                "institution_id": institution_id or current_user.get("institution_id"),
                "accreditor": accreditor or "SACSCOC"  # Default to SACSCOC
            }
            
            uploaded_files.append(file_info)
            
            if auto_analyze:
                # Start background analysis
                background_tasks.add_task(
                    analyze_document,
                    file_id,
                    file_path,
                    file_info
                )
        
        return {
            "success": True,
            "message": f"Uploaded {len(uploaded_files)} files for analysis",
            "files": uploaded_files
        }
        
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/status/{file_id}")
async def get_analysis_status(
    file_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Get the status of document analysis"""
    try:
        # In a real implementation, this would check a database
        # For now, return mock status
        status_file = Path(f"/tmp/evidence_analysis_{file_id}.json")
        
        if status_file.exists():
            with open(status_file, "r") as f:
                analysis_result = json.load(f)
            return analysis_result
        else:
            return {
                "file_id": file_id,
                "status": "processing",
                "progress": 50,
                "message": "Analyzing document..."
            }
            
    except Exception as e:
        logger.error(f"Status check error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/analyze/{file_id}")
async def trigger_analysis(
    file_id: str,
    accreditor: Optional[str] = Body(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: Dict = Depends(get_current_user)
):
    """Manually trigger analysis for an uploaded document"""
    try:
        # Find the file
        file_pattern = Path(f"/tmp/evidence_uploads/{file_id}_*")
        matching_files = list(Path("/tmp/evidence_uploads").glob(f"{file_id}_*"))
        
        if not matching_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_path = matching_files[0]
        file_info = {
            "id": file_id,
            "filename": file_path.name.replace(f"{file_id}_", ""),
            "accreditor": accreditor or "SACSCOC",
            "user_id": current_user.get("user_id")
        }
        
        # Start analysis
        background_tasks.add_task(
            analyze_document,
            file_id,
            file_path,
            file_info
        )
        
        return {
            "success": True,
            "message": "Analysis started",
            "file_id": file_id
        }
        
    except Exception as e:
        logger.error(f"Trigger analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/results")
async def get_all_results(
    current_user: Dict = Depends(get_current_user)
):
    """Get all analysis results for the current user"""
    try:
        results = []
        result_dir = Path("/tmp")
        
        # Find all analysis results for this user
        for result_file in result_dir.glob("evidence_analysis_*.json"):
            try:
                with open(result_file, "r") as f:
                    result = json.load(f)
                    if result.get("user_id") == current_user.get("user_id"):
                        results.append(result)
            except:
                continue
        
        # Sort by analysis date
        results.sort(key=lambda x: x.get("analyzed_at", ""), reverse=True)
        
        return {
            "success": True,
            "count": len(results),
            "results": results
        }
        
    except Exception as e:
        logger.error(f"Get results error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def analyze_document(file_id: str, file_path: Path, file_info: Dict[str, Any]):
    """Background task to analyze a document"""
    try:
        logger.info(f"Starting analysis for {file_id}")
        
        # Update status
        save_analysis_status(file_id, "processing", 10, "Reading document...", file_info)
        
        # Read document content
        processor = await get_evidence_processor()
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif file_path.suffix.lower() in ['.doc', '.docx']:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            # Plain text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        # Update status
        save_analysis_status(file_id, "processing", 30, "Extracting evidence...", file_info)
        
        # Create evidence document
        evidence_doc = EvidenceDocument(
            id=file_id,
            text=text,
            title=file_info.get("filename", "Document"),
            type="general",
            source=file_info.get("filename", "Unknown"),
            metadata={"file_info": file_info}
        )
        
        # Update status
        save_analysis_status(file_id, "processing", 50, "Mapping to standards...", file_info)
        
        # Get evidence mapper and map to standards
        mapper = get_evidence_mapper()
        accreditor = file_info.get("accreditor", "SACSCOC")
        
        # Filter corpus to specific accreditor if provided
        mapping_results = mapper.map_evidence(
            evidence_doc, 
            top_k=10,  # Get top 10 matches
            min_confidence=0.3  # Minimum confidence threshold
        )
        
        # Filter results by accreditor
        filtered_results = [
            r for r in mapping_results 
            if r.accreditor == accreditor or accreditor == "ALL"
        ]
        
        # Update status
        save_analysis_status(file_id, "processing", 80, "Calculating trust scores...", file_info)
        
        # Calculate overall trust score
        if filtered_results:
            avg_confidence = sum(r.confidence for r in filtered_results) / len(filtered_results)
            trust_score = min(95, int(avg_confidence * 100))
        else:
            trust_score = 0
        
        # Prepare final results
        analysis_result = {
            "file_id": file_id,
            "filename": file_info.get("filename"),
            "status": "completed",
            "progress": 100,
            "message": "Analysis complete",
            "user_id": file_info.get("user_id"),
            "analyzed_at": datetime.utcnow().isoformat(),
            "accreditor": accreditor,
            "trust_score": trust_score,
            "total_standards_mapped": len(filtered_results),
            "text_length": len(text),
            "mappings": [
                {
                    "standard_id": result.standard_id,
                    "standard_title": result.standard_title,
                    "confidence": round(result.confidence, 3),
                    "match_type": result.match_type,
                    "explanation": result.explanation,
                    "rationale_spans": [
                        {
                            "text": span.text[:200] + "..." if len(span.text) > 200 else span.text,
                            "start": span.start,
                            "end": span.end
                        }
                        for span in result.rationale_spans[:3]  # Limit to 3 spans
                    ]
                }
                for result in filtered_results[:10]  # Limit to top 10 results
            ]
        }
        
        # Save results
        save_analysis_status(file_id, "completed", 100, "Analysis complete", file_info, analysis_result)
        
        logger.info(f"Completed analysis for {file_id}: {len(filtered_results)} standards mapped")
        
    except Exception as e:
        logger.error(f"Analysis error for {file_id}: {e}")
        save_analysis_status(file_id, "failed", 0, str(e), file_info)

def save_analysis_status(file_id: str, status: str, progress: int, message: str, 
                        file_info: Dict[str, Any], result: Optional[Dict] = None):
    """Save analysis status to temporary storage"""
    try:
        status_data = {
            "file_id": file_id,
            "status": status,
            "progress": progress,
            "message": message,
            "updated_at": datetime.utcnow().isoformat()
        }
        
        if file_info:
            status_data.update({
                "filename": file_info.get("filename"),
                "user_id": file_info.get("user_id"),
                "institution_id": file_info.get("institution_id"),
                "accreditor": file_info.get("accreditor")
            })
        
        if result:
            status_data.update(result)
        
        # Save to temp file (in production, use database)
        status_file = Path(f"/tmp/evidence_analysis_{file_id}.json")
        with open(status_file, "w") as f:
            json.dump(status_data, f, indent=2)
            
    except Exception as e:
        logger.error(f"Failed to save status: {e}")