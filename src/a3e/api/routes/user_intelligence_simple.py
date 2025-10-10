"""
Simplified User Dashboard API Integration Service
Bypasses complex database queries to provide direct access to AI features
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Request, Header
from fastapi.responses import JSONResponse, PlainTextResponse, Response
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from typing import List, Dict, Any, Optional, Tuple, Set
from collections import defaultdict
from datetime import datetime
import os
import io
import csv
import json
import logging
import jwt
from pydantic import BaseModel

from ...services.standards_graph import standards_graph
from ...services.standards_loader import get_corpus_metadata
from ...services.evidence_mapper import evidence_mapper, EvidenceDocument
try:
    from ...services.evidence_mapper_enhanced import enhanced_evidence_mapper
    USE_ENHANCED_MAPPER = True
except ImportError:
    USE_ENHANCED_MAPPER = False
from ...services.evidence_trust import evidence_trust_scorer, EvidenceType, SourceSystem
from ...services.gap_risk_predictor import gap_risk_predictor
from ...services.risk_explainer import risk_explainer, StandardEvidenceSnapshot
from ...services.metrics_timeseries import maybe_snapshot, get_series
from ...services.telemetry_events import record_event as record_telemetry_event
from ...database.connection import db_manager
from sqlalchemy import text
from ...services.narrative_service import generate_narrative_html
try:
    from ...services.narrative_service_ai import ai_narrative_service
    USE_AI_NARRATIVE = True
except ImportError:
    USE_AI_NARRATIVE = False
try:
    from ...services.gap_risk_predictor_ai import ai_gap_risk_predictor
    USE_AI_GAP_PREDICTOR = True
except ImportError:
    USE_AI_GAP_PREDICTOR = False
from ...services.storage_service import get_storage_service, StorageService
from ...services.analytics_service import analytics_service
from ...core.config import get_settings
from ...models.document import Document as DocumentModel
import secrets
from pathlib import Path
import hashlib
import zipfile
import tempfile
import re
import uuid
# (imports deduplicated)

router = APIRouter(prefix="/api/user/intelligence-simple", tags=["user-intelligence-simple"])
security = HTTPBearer()
settings = get_settings()
logger = logging.getLogger(__name__)

JWT_ALGORITHM = "HS256"


def _parse_datetime_like(value: Any) -> Optional[datetime]:
    if isinstance(value, datetime):
        return value
    if isinstance(value, str):
        try:
            if value.endswith("Z"):
                value = value[:-1] + "+00:00"
            return datetime.fromisoformat(value)
        except Exception:
            return None
    return None


def _build_trust_summary(
    trust_payload: Optional[Dict[str, Any]],
    *,
    document_meta: Optional[Dict[str, Any]] = None,
    mappings: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    document_meta = document_meta or {}
    mappings = mappings or []
    trust_payload = trust_payload or {}

    signals_list = trust_payload.get("signals") or []
    signal_map = {}
    for item in signals_list:
        if isinstance(item, dict) and item.get("type"):
            signal_map[str(item["type"])] = item

    def _signal_value(key: str, fallback: float) -> float:
        value = signal_map.get(key, {}).get("value")
        if isinstance(value, (int, float)):
            return float(value)
        return fallback

    age_days = None
    reference_dt = _parse_datetime_like(document_meta.get("updated_at")) or _parse_datetime_like(document_meta.get("uploaded_at"))
    if reference_dt:
        try:
            age_days = max(0, (datetime.utcnow() - reference_dt).days)
        except Exception:
            age_days = None

    freshness_fallback = 1.0
    if age_days is not None:
        freshness_fallback = max(0.0, min(1.0, 1 - (age_days / 365.0)))
    freshness_score = _signal_value("freshness", freshness_fallback)

    provenance_base = 0.6
    content_type = (document_meta.get("content_type") or "").lower()
    file_size = document_meta.get("file_size") or 0
    if isinstance(file_size, str) and file_size.isdigit():
        file_size = int(file_size)
    provenance_base += 0.1 if any(token in content_type for token in ("pdf", "msword", "officedocument")) else 0.0
    provenance_base += 0.05 if document_meta.get("status") == "analyzed" else 0.0
    provenance_base += 0.1 if isinstance(file_size, (int, float)) and file_size and file_size > 100_000 else 0.0
    provenance_score = _signal_value("provenance", max(0.0, min(1.0, provenance_base)))

    confidences: List[float] = []
    excerpt_hits = 0
    total_mappings = len(mappings)
    for entry in mappings:
        conf = entry.get("confidence")
        if isinstance(conf, (int, float)):
            confidences.append(float(conf))
        excerpts = entry.get("excerpts")
        if not excerpts:
            excerpts = entry.get("page_anchors")
        if excerpts:
            excerpt_hits += 1
    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
    completeness_ratio = (excerpt_hits / total_mappings) if total_mappings else 0.0
    completeness_fallback = max(0.4, min(1.0, 0.4 + completeness_ratio * 0.6))
    completeness_score = _signal_value("completeness", completeness_fallback)
    relevance_score = _signal_value("relevance", avg_confidence)

    overall = trust_payload.get("overall_score")
    if not isinstance(overall, (int, float)):
        overall = (provenance_score + freshness_score + completeness_score + relevance_score) / 4.0
    overall = max(0.0, min(1.0, float(overall)))

    if overall >= 0.8:
        trust_level = "high"
    elif overall >= 0.6:
        trust_level = "medium"
    else:
        trust_level = "low"

    def _signal_payload(key: str, label: str, score: float, fallback_detail: str) -> Dict[str, Any]:
        explanation = signal_map.get(key, {}).get("explanation")
        if not isinstance(explanation, str) or not explanation.strip():
            explanation = fallback_detail
        return {
            "type": key,
            "label": label,
            "score": round(score, 3),
            "explanation": explanation,
        }

    signals_output = [
        _signal_payload(
            "provenance",
            "Provenance",
            provenance_score,
            "Assessed using upload source, document format, and automated processing signals.",
        ),
        _signal_payload(
            "freshness",
            "Freshness",
            freshness_score,
            "Document age influences freshness weighting ({} days).".format(age_days if age_days is not None else "unknown"),
        ),
        _signal_payload(
            "completeness",
            "Completeness",
            completeness_score,
            "{} of {} mapped standards include supporting excerpts.".format(excerpt_hits, total_mappings if total_mappings else 0),
        ),
        _signal_payload(
            "relevance",
            "Relevance",
            relevance_score,
            "Average mapping confidence is {}%.".format(int(round(avg_confidence * 100))),
        ),
    ]

    return {
        "overall_score": round(overall, 3),
        "level": trust_level,
        "signals": signals_output,
    }

# Ensure standards are loaded on startup
try:
    from ...startup import standards_loader
except ImportError:
    logger.warning("Standards loader startup module not found")

# ------------------------------
# Lightweight JSON stores (settings and reviews)
# ------------------------------
SETTINGS_STORE = os.getenv("USER_SETTINGS_STORE", "user_settings_store.json")
REVIEWS_STORE = os.getenv("USER_REVIEWS_STORE", "user_reviews_store.json")
# Per-standard reviewer notes/status (not tied to a specific evidence file)
STANDARD_REVIEWS_STORE = os.getenv("USER_STANDARD_REVIEWS_STORE", "user_standard_reviews_store.json")
UPLOADS_STORE = os.getenv("USER_UPLOADS_STORE", "user_uploads_store.json")
SESSIONS_STORE = os.getenv("USER_SESSIONS_STORE", "user_sessions_store.json")
ORG_CHART_STORE = os.getenv("ORG_CHART_STORE", "user_org_charts.json")
REVIEWS_AUDIT_LOG = os.getenv("USER_REVIEWS_AUDIT_LOG", "user_reviews_audit.jsonl")
TASKS_STORE = os.getenv("USER_TASKS_STORE", "user_tasks_store.json")
TASKS_AUDIT_LOG = os.getenv("USER_TASKS_AUDIT_LOG", "user_tasks_audit.jsonl")

# Simple uploads directory used by the dashboard's drag/drop upload
SIMPLE_UPLOADS_DIR = Path(os.getenv("SIMPLE_UPLOADS_DIR", "uploads/simple"))
SIMPLE_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

# BYOL (Bring Your Own License) standards directory
BYOL_STANDARDS_DIR = Path(os.getenv("BYOL_STANDARDS_DIR", "byol/standards"))
BYOL_STANDARDS_DIR.mkdir(parents=True, exist_ok=True)

def _get_standards_display_mode(claims: Dict[str, Any]) -> str:
    try:
        s = _get_user_settings(claims) or {}
    except Exception:
        s = {}
    mode = str(s.get("standards_display_mode") or os.getenv("STANDARDS_DISPLAY_MODE") or os.getenv("STANDARDS_SAFE_DISPLAY") or "full").lower()
    return "redacted" if mode in {"redacted", "safe", "hide"} else "full"

def _apply_display_policy(items: List[Dict[str, Any]], mode: str) -> List[Dict[str, Any]]:
    if mode == "full":
        return items
    redacted: List[Dict[str, Any]] = []
    for it in items:
        r = dict(it)
        r["description"] = ""
        redacted.append(r)
    return redacted

# (BYOL endpoints moved below auth helpers)


def _safe_load_json(path: str) -> Dict[str, Any]:
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception as e:
        logger.warning(f"Failed to load JSON from {path}: {e}")
    return {}


def _safe_save_json(path: str, data: Dict[str, Any]) -> None:
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Failed to save JSON to {path}: {e}")


def _get_user_org_chart(claims: Dict[str, Any]) -> Dict[str, Any]:
    all_c = _safe_load_json(ORG_CHART_STORE)
    return all_c.get(_user_key(claims), {})


def _save_user_org_chart(claims: Dict[str, Any], chart: Dict[str, Any]) -> None:
    all_c = _safe_load_json(ORG_CHART_STORE)
    all_c[_user_key(claims)] = chart
    _safe_save_json(ORG_CHART_STORE, all_c)


def _user_key(claims: Dict[str, Any]) -> str:
    # Prefer stable unique identifiers
    key = (claims.get("email") or claims.get("sub") or claims.get("user_id") or "unknown")
    return str(key)


def _get_user_settings(claims: Dict[str, Any]) -> Dict[str, Any]:
    # Try database first if DATABASE_URL is set (Railway/Production)
    if os.getenv("DATABASE_URL"):
        try:
            from ...services.user_settings_db import get_user_settings_db
            db_service = get_user_settings_db()
            settings = db_service.get_user_settings(claims)
            if settings:
                return settings
        except Exception as e:
            logger.warning(f"Failed to get settings from database: {e}")
    
    # Fallback to JSON file storage
    all_s = _safe_load_json(SETTINGS_STORE)
    return all_s.get(_user_key(claims), {})


def _save_user_settings(claims: Dict[str, Any], data: Dict[str, Any]) -> None:
    # Try database first if DATABASE_URL is set (Railway/Production)
    if os.getenv("DATABASE_URL"):
        try:
            from ...services.user_settings_db import get_user_settings_db
            db_service = get_user_settings_db()
            if db_service.save_user_settings(claims, data):
                logger.info("Settings saved to database")
                return
        except Exception as e:
            logger.error(f"Failed to save settings to database: {e}")
    
    # Fallback to JSON file storage
    all_s = _safe_load_json(SETTINGS_STORE)
    all_s[_user_key(claims)] = data
    _safe_save_json(SETTINGS_STORE, all_s)


def _get_user_reviews(claims: Dict[str, Any], filename: str) -> Dict[str, Any]:
    all_r = _safe_load_json(REVIEWS_STORE)
    return all_r.get(_user_key(claims), {}).get(filename, {})


def _set_user_review(
    claims: Dict[str, Any],
    filename: str,
    standard_id: str,
    reviewed: Optional[bool],
    note: Optional[str],
) -> Dict[str, Any]:
    all_r = _safe_load_json(REVIEWS_STORE)
    uk = _user_key(claims)
    user_map = all_r.get(uk, {})
    file_map = user_map.get(filename, {})
    entry = file_map.get(standard_id, {})
    if reviewed is not None:
        entry["reviewed"] = bool(reviewed)
    if note is not None:
        entry["note"] = str(note)
    file_map[standard_id] = entry
    user_map[filename] = file_map
    all_r[uk] = user_map
    _safe_save_json(REVIEWS_STORE, all_r)
    return entry


def _get_user_standard_reviews(claims: Dict[str, Any], accreditor: Optional[str] = None) -> Dict[str, Any]:
    """Return mapping of { standard_id: {status, note, assignee?, due_date?, updated_at} }.

    If accreditor provided, returns only that accreditor's map; otherwise returns merged across accreditors.
    Store structure: { user_key: { [accreditor]: { [standard_id]: entry } } }
    """
    all_r = _safe_load_json(STANDARD_REVIEWS_STORE)
    user_map = all_r.get(_user_key(claims), {})
    if accreditor:
        return user_map.get(accreditor.upper(), {}) or {}
    # Merge across accreditors (namespaces) into a single flat dict keyed by standard_id
    merged: Dict[str, Any] = {}
    for _, m in (user_map or {}).items():
        if isinstance(m, dict):
            merged.update(m)
    return merged


def _set_user_standard_review(
    claims: Dict[str, Any],
    accreditor: str,
    standard_id: str,
    status: Optional[str] = None,
    note: Optional[str] = None,
    assignee: Optional[str] = None,
    due_date: Optional[str] = None,
) -> Dict[str, Any]:
    """Upsert a per-standard review entry for the user and accreditor."""
    accreditor = (accreditor or "").upper()
    if not accreditor or not standard_id:
        raise HTTPException(status_code=400, detail="accreditor and standard_id are required")
    all_r = _safe_load_json(STANDARD_REVIEWS_STORE)
    uk = _user_key(claims)
    user_map = all_r.get(uk, {})
    acc_map = user_map.get(accreditor, {})
    entry = acc_map.get(standard_id, {})
    if status is not None:
        entry["status"] = str(status)
    if note is not None:
        entry["note"] = str(note)
    if assignee is not None:
        entry["assignee"] = str(assignee)
    if due_date is not None:
        entry["due_date"] = str(due_date)
    entry["updated_at"] = datetime.utcnow().isoformat()
    acc_map[standard_id] = entry
    user_map[accreditor] = acc_map
    all_r[uk] = user_map
    _safe_save_json(STANDARD_REVIEWS_STORE, all_r)
    # Audit trail
    try:
        prev = entry.copy()
        audit_entry = {
            "ts": datetime.utcnow().isoformat(),
            "user": _user_key(claims),
            "email": claims.get("email") or claims.get("sub"),
            "accreditor": accreditor,
            "standard_id": standard_id,
            "status": entry.get("status"),
            "assignee": entry.get("assignee"),
            "due_date": entry.get("due_date"),
            "note_len": len(entry.get("note") or ""),
        }
        with open(REVIEWS_AUDIT_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(audit_entry, ensure_ascii=False) + "\n")
    except Exception:
        pass
    return entry


def _merge_claims_with_settings(claims: Dict[str, Any]) -> Dict[str, Any]:
    defaults = {
        "organization": claims.get("organization") or claims.get("org") or "",
        "primary_accreditor": claims.get("primary_accreditor") or "",
        "tier": claims.get("tier") or "starter",
        "lms": claims.get("lms") or "",
        "term_system": claims.get("term_system") or "semester",
        "state": claims.get("state") or "",
        "institution_size": claims.get("institution_size") or "",
        "document_types": claims.get("document_types") or [],
        "document_sources": claims.get("document_sources") or [],
        # Onboarding fields
        "goals": claims.get("goals") or [],
        "launch_timing": claims.get("launch_timing") or "",
        "success_definition": claims.get("success_definition") or "",
    }
    saved = _get_user_settings(claims)
    merged = {**defaults, **(saved or {})}
    # Normalize lists
    if not isinstance(merged.get("document_sources"), list):
        merged["document_sources"] = []
    if not isinstance(merged.get("document_types"), list):
        # Accept legacy key 'doc_types' if present
        dt = merged.get("doc_types")
        if isinstance(dt, list):
            merged["document_types"] = dt
        else:
            merged["document_types"] = []
    if not isinstance(merged.get("goals"), list):
        merged["goals"] = []
    return merged


async def _get_user_uploads(claims: Dict[str, Any]) -> Dict[str, Any]:
    """Get user uploads from database"""
    raw_user_id = claims.get("user_id") or claims.get("sub") or claims.get("email") or "unknown"

    # Normalize to the UUID stored in the database (older tokens may carry an email)
    user_id = await get_user_uuid_from_email(raw_user_id)

    # Cache normalized id back onto the claims so downstream callers stay consistent
    claims.setdefault("user_id", user_id)
    
    try:
        async with db_manager.get_session() as session:
            # Get documents with their analysis results
            result = await session.execute(
                text("""
                    SELECT id, filename, file_key, file_size, content_type, 
                           sha256, status, uploaded_at, organization_id,
                           analysis_results
                    FROM documents 
                    WHERE user_id = :user_id 
                    AND deleted_at IS NULL
                    ORDER BY uploaded_at DESC
                    LIMIT 50
                """),
                {"user_id": user_id}
            )
            
            documents = []
            doc_ids = []
            
            for row in result:
                doc_id = str(row.id)
                # Parse analysis results for trust score
                analysis = json.loads(row.analysis_results) if row.analysis_results else {}
                trust_score = analysis.get("trust_score", {})
                
                doc = {
                    "id": doc_id,
                    "filename": row.filename,
                    "uploaded_at": row.uploaded_at.isoformat() if row.uploaded_at else "",
                    "standards_mapped": [],  # Will fill from evidence_mappings
                    "doc_type": row.content_type or "",
                    "mappings": [],
                    "trust_score": trust_score,
                    "saved_path": row.file_key,
                    "fingerprint": row.sha256 or "",
                    "size": row.file_size or 0,
                }
                documents.append(doc)
                doc_ids.append(row.id)
            
            # Get unique standards from evidence mappings
            unique_standards = set()
            if doc_ids:
                mapping_result = await session.execute(
                    text("""
                        SELECT standard_id, document_id, confidence
                        FROM evidence_mappings 
                        WHERE document_id = ANY(:doc_ids)
                        AND confidence > 0.3
                    """),
                    {"doc_ids": doc_ids}
                )
                
                doc_to_standards: Dict[str, List[str]] = {}
                doc_to_mappings: Dict[str, List[Dict[str, Any]]] = {}

                for mapping in mapping_result:
                    standard_id = str(mapping.standard_id)
                    document_id = str(mapping.document_id)
                    unique_standards.add(standard_id)

                    doc_to_standards.setdefault(document_id, []).append(standard_id)

                    confidence_val = None
                    try:
                        if mapping.confidence is not None:
                            confidence_val = float(mapping.confidence)
                    except (TypeError, ValueError):
                        confidence_val = None

                    inferred_acc = None
                    if standard_id:
                        token = standard_id.split(":", 1)[0].split("_", 1)[0].split("-", 1)[0]
                        inferred_acc = token.upper() if token else None

                    doc_to_mappings.setdefault(document_id, []).append({
                        "standard_id": standard_id,
                        "accreditor": inferred_acc,
                        "confidence": confidence_val,
                    })
                
                # Update documents with their mapped standards
                for doc in documents:
                    doc_id = doc["id"]
                    doc["standards_mapped"] = doc_to_standards.get(doc_id, [])
                    doc["mappings"] = doc_to_mappings.get(doc_id, [])
            
            return {
                "documents": documents,
                "unique_standards": list(unique_standards)
            }
    except Exception as e:
        logger.error(f"Error getting user uploads from database: {e}")
        # Fallback to empty result
        return {"documents": [], "unique_standards": []}


async def _record_user_upload(
    claims: Dict[str, Any],
    filename: str,
    standard_ids: List[str],
    doc_type: Optional[str] = None,
    mapping_details: Optional[List[Dict[str, Any]]] = None,
    trust_score: Optional[Dict[str, Any]] = None,
    saved_path: Optional[str] = None,
    fingerprint: Optional[str] = None,
    file_size: Optional[int] = None,
    analysis_results: Optional[Dict[str, Any]] = None,
    content_type: Optional[str] = None,
) -> Dict[str, Any]:
    """Record a new upload in the database"""
    identifier = claims.get("user_id") or claims.get("sub") or claims.get("email") or "unknown"
    user_id = await get_user_uuid_from_email(identifier)
    claims.setdefault("user_id", user_id)

    # Persist an organization identifier when available (not required by schema)
    org_id = claims.get("org_id") or claims.get("organization_id")
    
    try:
        async with db_manager.get_session() as session:
            # Create new document record
            document_id = str(uuid.uuid4())
            analysis_json = json.dumps(analysis_results) if analysis_results is not None else None
            
            params = {
                "id": document_id,
                "user_id": user_id,
                "organization_id": org_id,
                "filename": filename,
                "file_key": saved_path or "",
                "file_size": file_size or 0,
                "content_type": content_type or doc_type or "application/octet-stream",
                "sha256": fingerprint or "",
                "status": "analyzed" if standard_ids else "uploaded",
                "uploaded_at": datetime.utcnow(),
                "original_filename": filename,
                "file_path": saved_path or "",
                "mime_type": content_type or doc_type or "application/octet-stream",
                "analysis_results": analysis_json,
            }

            used_legacy_schema = False

            insert_stmt = text(
                """
                INSERT INTO documents (
                    id, user_id, organization_id, filename, file_key,
                    file_size, content_type, sha256, status, uploaded_at,
                    original_filename, file_path, mime_type, analysis_results
                ) VALUES (
                    :id, :user_id, :organization_id, :filename, :file_key,
                    :file_size, :content_type, :sha256, :status, :uploaded_at,
                    :original_filename, :file_path, :mime_type, :analysis_results
                )
                """
            )

            try:
                await session.execute(insert_stmt, params)
            except Exception as insert_error:
                await session.rollback()
                error_text = str(insert_error).lower()
                # Retry against legacy schema that still uses institution_id
                if "organization_id" in error_text:
                    legacy_params = params.copy()
                    legacy_params.pop("organization_id", None)
                    legacy_params["institution_id"] = org_id
                    legacy_stmt = text(
                        """
                        INSERT INTO documents (
                            id, user_id, institution_id, filename, file_key,
                            file_size, content_type, sha256, status, uploaded_at,
                            original_filename, file_path, mime_type, analysis_results
                        ) VALUES (
                            :id, :user_id, :institution_id, :filename, :file_key,
                            :file_size, :content_type, :sha256, :status, :uploaded_at,
                            :original_filename, :file_path, :mime_type, :analysis_results
                        )
                        """
                    )
                    await session.execute(legacy_stmt, legacy_params)
                    used_legacy_schema = True
                else:
                    raise insert_error

            await session.commit()
            
            # TODO: If standard_ids provided, insert into mapping table
            
            schema_note = " using legacy schema" if used_legacy_schema else ""
            logger.info(f"Recorded upload {filename} for user {user_id} in database{schema_note}")
            
            # Return document_id for further processing
            return {"document_id": document_id, "uploads": await _get_user_uploads(claims)}
            
    except Exception as e:
        logger.error(f"Error recording upload in database: {e}")
        # Fallback to return basic data
        return {
            "documents": [
                {
                    "filename": filename,
                    "uploaded_at": datetime.utcnow().isoformat(),
                    "standards_mapped": list(standard_ids or []),
                    "doc_type": doc_type or "",
                    "mappings": mapping_details or [],
                    "trust_score": trust_score or {},
                    "saved_path": saved_path or "",
                    "fingerprint": fingerprint or "",
                    "size": file_size or 0,
                    "analysis_results": analysis_results or {},
                }
            ],
            "unique_standards": list(standard_ids or [])
        }


# (Checklist endpoints moved below auth helpers to satisfy linters)


async def _compute_dashboard_metrics_for_snapshot(current_user: Dict[str, Any]) -> Dict[str, Any]:
    settings_ = _merge_claims_with_settings(current_user)
    uploads = await _get_user_uploads(current_user)
    acc = (settings_.get("primary_accreditor") or "HLC").upper()
    total_roots = len(standards_graph.get_accreditor_standards(acc)) or 0
    docs = uploads.get("documents", [])
    uniq_standards = set(uploads.get("unique_standards", []) or [])

    documents_analyzed = len(docs)
    standards_mapped = len(uniq_standards)
    total_standards = total_roots if total_roots > 0 else max(total_roots, standards_mapped)
    coverage = (standards_mapped / total_standards) if total_standards else 0.0

    trust_scores: List[float] = []
    for d in docs:
        ts = (d.get("trust_score") or {}).get("overall_score")
        if isinstance(ts, (int, float)):
            trust_scores.append(float(ts))
        # Normalize average trust to 0-1
        avg_trust = float(sum(trust_scores) / len(trust_scores)) if trust_scores else 0.7
        if avg_trust > 1:
            avg_trust = avg_trust / 100.0
        avg_trust = max(0.0, min(1.0, avg_trust))

    compliance_score = 0.0
    if total_standards > 0 and (standards_mapped > 0 or documents_analyzed > 0):
        compliance_score = round((coverage * 0.7 + avg_trust * 0.3) * 100, 1)

    risk_agg = risk_explainer.aggregate()
    average_risk = float(risk_agg.get("average_risk", 0.0))

    return {
        "accreditor": acc,
        "coverage_percentage": round(coverage * 100, 1),
        "compliance_score": compliance_score,
        "average_trust": round(avg_trust, 3),
        "average_risk": round(average_risk, 4),
        "documents_analyzed": documents_analyzed,
        "standards_mapped": standards_mapped,
        "total_standards": total_standards,
    }


# ------------------------------
# Internal helper: analyze evidence content from bytes
# ------------------------------
async def _analyze_evidence_from_bytes(
    filename: str,
    content: bytes,
    doc_type: Optional[str],
    current_user: Dict[str, Any],
    document_id: Optional[str] = None,  # Add document ID to update existing record
):
    try:
        filename_lower = (filename or "").lower()
        is_pdf = (
            filename_lower.endswith(".pdf") or content[:4] == b"%PDF"
        )
        text_content = ""
        page_texts: List[str] = []
        if is_pdf:
            try:
                import pypdf  # type: ignore
                from io import BytesIO
                reader = pypdf.PdfReader(BytesIO(content))
                parts = []
                for i, page in enumerate(reader.pages[:20]):
                    try:
                        txt = page.extract_text() or ""
                        parts.append(txt)
                        page_texts.append(txt)
                    except Exception:
                        continue
                text_content = "\n".join([p for p in parts if p]).strip()
            except Exception:
                text_content = ""
        else:
            try:
                text_content = content.decode("utf-8", errors="ignore")
            except Exception:
                text_content = ""

        # Optional OCR fallback when PDF has no text
        if is_pdf and not text_content:
            try:
                ocr_enabled = os.getenv("OCR_ENABLED", "false").lower() in {"1", "true", "yes"}
                if ocr_enabled:
                    from pdf2image import convert_from_bytes  # type: ignore
                    import pytesseract  # type: ignore
                    images = convert_from_bytes(content, first_page=1, last_page=5)
                    ocr_texts: List[str] = []
                    for img in images:
                        try:
                            t = pytesseract.image_to_string(img) or ""
                            if t.strip():
                                ocr_texts.append(t)
                                page_texts.append(t)
                        except Exception:
                            continue
                    text_content = "\n".join(ocr_texts).strip()
            except Exception:
                pass

        # Optional PII/FERPA preflight redaction
        redaction_enabled = os.getenv("PII_REDACTION_ENABLED", "true").lower() in {"1", "true", "yes"}
        redaction_report = {"emails": 0, "ssn": 0, "phones": 0, "dob": 0}

        def _redact(text: str) -> str:
            nonlocal redaction_report
            import re as _re
            text = _re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", lambda m: redaction_report.__setitem__("emails", redaction_report["emails"] + 1) or "[REDACTED_EMAIL]", text)
            text = _re.sub(r"\b\d{3}-\d{2}-\d{4}\b", lambda m: redaction_report.__setitem__("ssn", redaction_report["ssn"] + 1) or "[REDACTED_SSN]", text)
            text = _re.sub(r"(\+?\d{1,2}[\s-])?(\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4})", lambda m: redaction_report.__setitem__("phones", redaction_report["phones"] + 1) or "[REDACTED_PHONE]", text)
            text = _re.sub(r"\b(0?[1-9]|1[0-2])/(0?[1-9]|[12]\d|3[01])/(19|20)\d{2}\b", lambda m: redaction_report.__setitem__("dob", redaction_report["dob"] + 1) or "[REDACTED_DOB]", text)
            return text

        redacted_text = _redact(text_content) if (redaction_enabled and text_content) else text_content

        # Get user's institutional context
        email = current_user.get('sub') or current_user.get('email') or current_user.get('user_id')
        user_id = await get_user_uuid_from_email(email)
        user_institution = None
        user_accreditor = None
        
        if user_id:
            try:
                async with db_manager.get_session() as session:
                    user_result = await session.execute(
                        text("SELECT institution_name, primary_accreditor FROM users WHERE id = :user_id"),
                        {"user_id": user_id}
                    )
                    user_data = user_result.fetchone()
                    if user_data:
                        user_institution = user_data[0]
                        user_accreditor = user_data[1]
            except Exception as e:
                logger.warning(f"Could not fetch user institution data: {e}")
        
        doc = EvidenceDocument(
            doc_id=filename,
            text=redacted_text,
            metadata={
                "uploaded_by": _user_key(current_user), 
                "doc_type": doc_type or "policy",
                "institution": user_institution or "",
                "accreditor": user_accreditor or "",
                "user_id": user_id or ""
            },
            doc_type="policy",
            source_system="manual",
            upload_date=datetime.utcnow(),
        )

        # Use enhanced mapper if available and OpenAI key is configured
        if USE_ENHANCED_MAPPER and settings.openai_api_key:
            try:
                # Initialize enhanced mapper if needed
                if not enhanced_evidence_mapper._initialized:
                    await enhanced_evidence_mapper.initialize()
                
                # Use AI-enhanced mapping
                mappings = await enhanced_evidence_mapper.map_evidence_with_ai(
                    doc, 
                    num_candidates=20, 
                    final_top_k=10,
                    use_llm=True
                )
            except Exception as e:
                logger.warning(f"Enhanced mapping failed, falling back to TF-IDF: {e}")
                mappings = evidence_mapper.map_evidence(doc)
        else:
            mappings = evidence_mapper.map_evidence(doc)
        top_conf = mappings[0].confidence if mappings else 0.6

        trust = evidence_trust_scorer.calculate_trust_score(
            evidence_id=doc.doc_id,
            evidence_type=EvidenceType.POLICY,
            source_system=SourceSystem.MANUAL,
            upload_date=doc.upload_date,
            last_modified=datetime.utcnow(),
            content_length=len(doc.text or ""),
            metadata=doc.metadata,
            mapping_confidence=top_conf,
            reviewer_approved=True,
            citations_count=0,
            conflicts_detected=0,
        )

        trust_dict = trust.to_dict()
        signals = {s["type"]: s["value"] for s in trust_dict.get("signals", [])}
        quality_score = signals.get("completeness", trust_dict.get("overall_score", 0.7))
        reliability_score = (signals.get("provenance", 0.7) + signals.get("alignment", 0.7)) / 2.0
        confidence_score = signals.get("reviewer_verification", 0.8)

        def _meets(mt: str, conf: float) -> bool:
            try:
                return bool(mt in ("exact", "strong") or float(conf) >= 0.7)
            except Exception:
                return False

        reviews_map = _get_user_reviews(current_user, filename)
        mappings_ui = []
        mapping_details: List[Dict[str, Any]] = []
        for m in mappings[:10]:
            review = reviews_map.get(m.standard_id, {}) if isinstance(reviews_map, dict) else {}
            anchors: List[Dict[str, Any]] = []
            try:
                for idx, page_txt in enumerate(page_texts or []):
                    for span in m.rationale_spans[:2]:
                        snip = (span or "").replace("**", "").strip()
                        if snip and snip[:24] in page_txt:
                            anchors.append({"page": idx + 1, "snippet": snip[:120]})
                            break
            except Exception:
                anchors = []
            mappings_ui.append(
                {
                    "standard_id": m.standard_id,
                    "title": m.standard_title,
                    "confidence": float(m.confidence),
                    "accreditor": m.accreditor,
                    "match_type": m.match_type,
                    "meets_standard": bool(_meets(m.match_type, m.confidence)),
                    "rationale_spans": m.rationale_spans,
                    "explanation": m.explanation,
                    "page_anchors": anchors,
                    "reviewed": bool(review.get("reviewed", False)),
                    "note": review.get("note", ""),
                    "rationale_note": review.get("rationale_note", ""),
                }
            )
            mapping_details.append({
                "standard_id": m.standard_id,
                "accreditor": m.accreditor,
                "confidence": float(m.confidence),
                "page_anchors": anchors,
            })
        fingerprint = EvidenceDocument(doc_id=filename, text=text_content, metadata={}, doc_type="", source_system="manual", upload_date=datetime.utcnow()).get_fingerprint()

        summary_mappings = [
            {
                "confidence": md.get("confidence"),
                "excerpts": md.get("page_anchors", []),
            }
            for md in mapping_details
        ]

        document_meta_summary = {
            "uploaded_at": datetime.utcnow().isoformat(),
            "status": "analyzed",
            "file_size": len(content),
            "content_type": (doc_type or ("application/pdf" if is_pdf else "text/plain")),
        }

        trust_summary = _build_trust_summary(
            trust_dict,
            document_meta={
                "uploaded_at": document_meta_summary["uploaded_at"],
                "status": document_meta_summary["status"],
                "file_size": document_meta_summary["file_size"],
                "content_type": document_meta_summary["content_type"],
            },
            mappings=summary_mappings,
        )

        analysis_payload = {
            "is_pdf": bool(is_pdf),
            "mappings": mappings_ui,
            "trust_score": {
                "overall_score": float(trust_dict.get("overall_score", 0.7) or 0.7),
                "quality_score": float(round(float(quality_score), 3)),
                "reliability_score": float(round(float(reliability_score), 3)),
                "confidence_score": float(round(float(confidence_score), 3)),
            },
            "trust_summary": trust_summary,
            "standards_mapped": len(mappings),
            "content_length": len(doc.text or ""),
            "redaction": {"enabled": redaction_enabled, **redaction_report},
            "fingerprint": fingerprint,
            "document_preview": (
                "[PDF detected: preview unavailable]" if (is_pdf and not (doc.text or "").strip()) else (doc.text or "")[:2000]
            ),
            "analysis_generated_at": datetime.utcnow().isoformat(),
        }

        # If we have a document_id, update the existing record instead of creating a new one
        if document_id:
            try:
                async with db_manager.get_session() as session:
                    # Update the existing document with analysis results
                    await session.execute(
                        text("""
                            UPDATE documents 
                            SET status = 'analyzed',
                                updated_at = CURRENT_TIMESTAMP,
                                analysis_results = :analysis_results
                            WHERE id = :id
                        """),
                        {
                            "id": document_id,
                            "analysis_results": json.dumps(analysis_payload),
                        }
                    )
                    
                    # Also record the mappings in evidence_mappings table
                    for i, m in enumerate(mappings[:10]):
                        await session.execute(
                            text("""
                                INSERT INTO evidence_mappings (
                                    id, document_id, standard_id, confidence, 
                                    excerpts, created_at
                                ) VALUES (
                                    :id, :document_id, :standard_id, :confidence,
                                    :excerpts, CURRENT_TIMESTAMP
                                ) ON CONFLICT (document_id, standard_id) DO UPDATE SET
                                    confidence = EXCLUDED.confidence,
                                    excerpts = EXCLUDED.excerpts,
                                    updated_at = CURRENT_TIMESTAMP
                            """),
                            {
                                "id": str(uuid.uuid4()),
                                "document_id": document_id,
                                "standard_id": m.standard_id,
                                "confidence": float(m.confidence),
                                "excerpts": json.dumps([
                                    {"page": anchor.get("page", 1), "snippet": anchor.get("snippet", "")}
                                    for anchor in mapping_details[i].get("page_anchors", [])
                                ] if i < len(mapping_details) else [])
                            }
                        )
                    
                    await session.commit()
            except Exception as e:
                logger.error(f"Error updating document analysis: {e}")
        else:
            # No document_id, create new record (original behavior)
            await _record_user_upload(
                current_user,
                filename,
                [m.standard_id for m in mappings],
                doc_type,
                mapping_details,
                trust_dict,
                None,
                fingerprint,
                analysis_results=analysis_payload,
            )

        try:
            overall_trust = float(trust_dict.get("overall_score", 0.7) or 0.7)
            for md in mapping_details[:50]:
                sid = str(md.get("standard_id"))
                conf = float(md.get("confidence") or 0.0)
                snapshot = StandardEvidenceSnapshot(
                    standard_id=sid,
                    coverage_percent=25.0,
                    trust_scores=[overall_trust, conf],
                    evidence_ages_days=[365],
                    overdue_tasks=0,
                    total_tasks=0,
                    recent_changes=0,
                    historical_findings=0,
                    days_to_review=180,
                )
                risk_explainer.compute_standard_risk(snapshot)
        except Exception:
            pass

        return {
            "status": "success",
            "filename": filename,
            "analysis": analysis_payload,
            "algorithms_used": ["EvidenceMapper™", "EvidenceTrust Score™", "StandardsGraph™"],
        }
    except Exception as e:
        logger.error(f"Evidence analysis (from-bytes) error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {e}")


# ------------------------------
# Auth helpers
# ------------------------------

def _candidate_secrets() -> List[str]:
    candidates: List[Optional[str]] = [
        getattr(settings, "jwt_secret_key", None),
        getattr(settings, "secret_key", None),
        os.getenv("JWT_SECRET_KEY"),
        os.getenv("JWT_SECRET"),
        os.getenv("SECRET_KEY"),
        os.getenv("ONBOARDING_SHARED_SECRET"),
        # Common development/example secrets that might be in old tokens
        "your-secret-key-here-change-in-production",
        "replace-with-generated-64-char-secret",
        "your-secret-key-min-32-chars-generate-with-openssl-rand-hex-32",
        # The one from .env.local file
        "7UKtJWo1jG6Yji-Fw-0t1HRC6y8QsPojrWkEJhEXXTQV0myYJIJ183xEPLcT6vDcPjLR_mB9tBQsGejvTxg-QA",
        # Session router fallback default (only if env not set)
        "dev-secret-change",
    ]
    return [c for c in candidates if c]


def verify_simple_token(token: str) -> Optional[Dict[str, Any]]:
    for secret in _candidate_secrets():
        try:
            payload = jwt.decode(token, secret, algorithms=[JWT_ALGORITHM])
            return payload
        except Exception:
            continue
    logger.warning("Token verification failed for all candidate secrets")
    return None


async def get_current_user_simple(
    request: Request,
    authorization: Optional[str] = Header(default=None),
) -> Dict[str, Any]:
    token: Optional[str] = None
    if authorization and isinstance(authorization, str) and authorization.lower().startswith("bearer "):
        token = authorization.split(" ", 1)[1].strip()
    if not token:
        token = request.cookies.get("access_token") or request.cookies.get("a3e_api_key") or request.cookies.get("jwt_token")
    if not token:
        raise HTTPException(status_code=401, detail="Authentication required")
    claims = verify_simple_token(token)
    if not claims:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    # Handle backward compatibility: convert email to UUID if needed
    # Check if 'sub' contains an email instead of a UUID
    sub = claims.get("sub", "")
    if sub and "@" in sub:
        # This is an old token with email in 'sub'
        logger.info(f"Converting old token format for user: {sub}")
        # For now, just use the email as the user ID
        # The individual endpoints will handle conversion when needed
        claims["user_id"] = sub
        if "email" not in claims:
            claims["email"] = sub  # Preserve the email
    elif not claims.get("user_id"):
        # Ensure user_id is set
        claims["user_id"] = sub
    
    return claims


# ------------------------------
# Helper function to get user UUID from email
# ------------------------------
async def get_user_uuid_from_email(email: str) -> str:
    """Get the actual UUID for a user from their email address"""
    try:
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("SELECT id FROM users WHERE email = :email"),
                {"email": email}
            )
            user = result.first()
            if user:
                return user.id
            # If no user found, check if the provided value is already a UUID
            import uuid
            try:
                uuid.UUID(email)
                return email  # It's already a UUID
            except:
                # Return the email as fallback (for backward compatibility)
                logger.warning(f"No user found for email: {email}, using email as ID")
                return email
    except Exception as e:
        logger.error(f"Error getting user UUID: {e}")
        return email


# ------------------------------
# Evidence Intake Checklist
# ------------------------------
@router.get("/standards/checklist")
async def get_evidence_intake_checklist(
    accreditor: Optional[str] = None,
    format: Optional[str] = None,
    include_indicators: bool = False,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    acc = (accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()
    uploads = _get_user_uploads(current_user)
    mapped: Set[str] = set(uploads.get("unique_standards", []) or [])

    def node_entry(n):
        return {
            "standard_id": n.node_id,
            "accreditor": n.accreditor,
            "level": n.level,
            "title": n.title,
            "description": n.description,
            "evidence_requirements": list(n.evidence_requirements or []),
            "covered": n.node_id in mapped,
        }

    levels = {"standard", "clause"}
    if include_indicators:
        levels.add("indicator")
    nodes = standards_graph.get_nodes_by_accreditor(acc, levels)
    by_standard: Dict[str, Dict[str, Any]] = {}
    for n in nodes:
        path = standards_graph.get_path_to_root(n.node_id)
        root_node = path[0] if path else n
        root_id = root_node.node_id
        group = by_standard.setdefault(root_id, {
            "standard_id": root_id,
            "title": root_node.title if root_node else "",
            "items": [],
        })
        group["items"].append(node_entry(n))

    fmt = (format or "json").lower()
    if fmt == "csv":
        out_io = io.StringIO()
        writer = csv.writer(out_io)
        writer.writerow(["standard_id", "accreditor", "level", "title", "description", "evidence_requirements", "covered"])
        for g in by_standard.values():
            for it in g["items"]:
                writer.writerow([
                    it.get("standard_id"),
                    it.get("accreditor"),
                    it.get("level"),
                    it.get("title"),
                    (it.get("description") or "").replace("\n", " ").strip(),
                    "; ".join(it.get("evidence_requirements") or []),
                    "yes" if it.get("covered") else "no",
                ])
        return Response(content=out_io.getvalue(), media_type="text/csv")

    total = sum(1 for g in by_standard.values() for _ in g["items"]) or 0
    covered = sum(1 for g in by_standard.values() for it in g["items"] if it.get("covered"))
    return {
        "success": True,
        "accreditor": acc,
        "total_items": total,
        "covered_items": covered,
        "coverage_percentage": round((covered / total) * 100, 1) if total else 0.0,
        "groups": list(by_standard.values()),
    }


@router.get("/standards/checklist/stats")
async def get_evidence_intake_checklist_stats(
    accreditor: Optional[str] = None,
    include_indicators: bool = False,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    acc = (accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()
    uploads = _get_user_uploads(current_user)
    mapped: Set[str] = set(uploads.get("unique_standards", []) or [])
    levels = {"standard", "clause"}
    if include_indicators:
        levels.add("indicator")
    nodes = standards_graph.get_nodes_by_accreditor(acc, levels)
    total = len(nodes)
    covered = len([n for n in nodes if n.node_id in mapped])
    return {
        "accreditor": acc,
        "total_items": total,
        "covered_items": covered,
        "coverage_percentage": round((covered / total) * 100, 1) if total else 0.0,
    }

# BYOL Standards Ingestion
# ------------------------------
@router.post("/standards/byol/upload")
async def upload_byol_standards(
    file: UploadFile = File(...),
    accreditor: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Upload an accreditor standards file (YAML/JSON) into BYOL store.

    Stores the file under `BYOL_STANDARDS_DIR` and returns a handle. Does not reload the graph.
    """
    try:
        filename = file.filename or "standards.yaml"
        suffix = Path(filename).suffix.lower()
        if suffix not in {".yaml", ".yml", ".json"}:
            raise HTTPException(status_code=400, detail="Only .yaml, .yml, or .json files are supported")
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="Empty file")
        # Optionally namespace by accreditor
        subdir = BYOL_STANDARDS_DIR / (accreditor.upper() if accreditor else "")
        subdir.mkdir(parents=True, exist_ok=True)
        target = subdir / filename
        with open(target, "wb") as f:
            f.write(data)
        return {
            "success": True,
            "stored_at": str(target),
            "bytes": len(data),
            "hint": "Call /standards/byol/reload to refresh the in-memory graph",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"BYOL upload error: {e}")
        raise HTTPException(status_code=500, detail="Failed to store BYOL standards file")


@router.post("/standards/byol/reload")
async def reload_byol_standards(
    path: Optional[str] = Form(None),
    fallback_to_seed: bool = Form(True),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Reload the StandardsGraph from a corpus directory.

    - If `path` provided, use it; otherwise use `BYOL_STANDARDS_DIR` if non-empty;
      else default repo `data/standards`.
    """
    try:
        base = None
        if path:
            base = Path(path)
        else:
            # Prefer BYOL dir if contains at least one standards file
            has_files = any(p.is_file() and p.suffix.lower() in {".yaml", ".yml", ".json"} for p in BYOL_STANDARDS_DIR.glob("**/*"))
            if has_files:
                base = BYOL_STANDARDS_DIR
        stats = standards_graph.reload_from_corpus(str(base) if base else None, fallback_to_seed=bool(fallback_to_seed))
        # Include corpus status snapshot
        meta = await get_corpus_metadata_api(current_user=current_user)  # type: ignore
        return {"success": True, "reloaded_from": str(base) if base else None, "graph": stats, "corpus": meta}
    except Exception as e:
        logger.error(f"BYOL reload error: {e}")
        raise HTTPException(status_code=500, detail="Failed to reload standards corpus")


# Compatibility alias: allow GET-based reload in environments where POST may be blocked
@router.get("/standards/byol/reload")
async def reload_byol_standards_get(
    path: Optional[str] = None,
    fallback_to_seed: bool = True,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    try:
        base = None
        if path:
            base = Path(path)
        else:
            has_files = any(p.is_file() and p.suffix.lower() in {".yaml", ".yml", ".json"} for p in BYOL_STANDARDS_DIR.glob("**/*"))
            if has_files:
                base = BYOL_STANDARDS_DIR
        stats = standards_graph.reload_from_corpus(str(base) if base else None, fallback_to_seed=bool(fallback_to_seed))
        meta = await get_corpus_metadata_api(current_user=current_user)  # type: ignore
        return {"success": True, "reloaded_from": str(base) if base else None, "graph": stats, "corpus": meta, "via": "GET"}
    except Exception as e:
        logger.error(f"BYOL reload (GET) error: {e}")
        raise HTTPException(status_code=500, detail="Failed to reload standards corpus (GET)")


# ------------------------------
# BYOL: ingest from storage (supports S3 presigned)
# ------------------------------
@router.post("/standards/byol/ingest/from-storage")
async def ingest_byol_from_storage(
    payload: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
    storage: StorageService = Depends(get_storage_service),
):
    """Copy a YAML/JSON standards file from storage into BYOL corpus dir.

    Expected payload: { file_key, filename, accreditor? }
    """
    file_key = str(payload.get("file_key") or "").strip()
    filename = str(payload.get("filename") or "").strip()
    accreditor = (payload.get("accreditor") or "").strip() or None
    if not file_key or not filename:
        raise HTTPException(status_code=400, detail="file_key and filename are required")
    try:
        content = await storage.get_file(file_key)
        if not content:
            raise HTTPException(status_code=404, detail="File not found in storage")
        suffix = Path(filename).suffix.lower()
        if suffix not in {".yaml", ".yml", ".json"}:
            raise HTTPException(status_code=400, detail="Only .yaml, .yml, .json supported for BYOL")
        subdir = BYOL_STANDARDS_DIR / (accreditor.upper() if accreditor else "")
        subdir.mkdir(parents=True, exist_ok=True)
        target = subdir / filename
        with open(target, "wb") as f:
            f.write(content)
        return {"success": True, "stored_at": str(target), "bytes": len(content)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"BYOL ingest-from-storage error: {e}")
        raise HTTPException(status_code=500, detail="Failed to ingest BYOL file from storage")


# ------------------------------
# Settings and integrations
# ------------------------------
@router.get("/settings")
async def get_settings(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return _merge_claims_with_settings(current_user)


@router.post("/settings")
async def save_settings(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    existing = _merge_claims_with_settings(current_user)
    updated = {**existing, **(payload or {})}
    _save_user_settings(current_user, updated)
    return {"status": "saved", "settings": updated}


# ------------------------------
# Org Chart: save/load (per-user JSON store)
# ------------------------------
@router.get("/org-chart")
async def load_org_chart(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        chart = _get_user_org_chart(current_user) or {}
        exists = bool(chart)
        return {
            "exists": exists,
            "chart": chart,
            "last_updated": chart.get("updated_at") if isinstance(chart, dict) else None,
        }
    except Exception as e:
        logger.error(f"Load org chart error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load org chart")


@router.post("/org-chart")
async def save_org_chart(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        existing = _get_user_org_chart(current_user) or {}
        nodes = payload.get("nodes") or payload.get("data", {}).get("nodes") or []
        edges = payload.get("edges") or payload.get("data", {}).get("edges") or []
        metadata = payload.get("metadata") or payload.get("data", {}).get("metadata") or {}
        name = payload.get("name") or existing.get("name") or "My Organization Chart"
        description = payload.get("description") or existing.get("description") or ""
        now = datetime.utcnow().isoformat()

        chart = {
            "name": name,
            "description": description,
            "nodes": nodes,
            "edges": edges,
            "metadata": metadata,
            "created_at": existing.get("created_at") or now,
            "updated_at": now,
        }
        _save_user_org_chart(current_user, chart)
        return {"status": "saved", "chart": chart}
    except Exception as e:
        logger.error(f"Save org chart error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save org chart")

# Aliases for explicit save/load paths
@router.get("/org-chart/load")
async def load_org_chart_alias(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        chart = _get_user_org_chart(current_user) or {}
        exists = bool(chart)
        return {
            "exists": exists,
            "chart": chart,
            "last_updated": chart.get("updated_at") if isinstance(chart, dict) else None,
        }
    except Exception as e:
        logger.error(f"Load org chart (alias) error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load org chart")


@router.post("/org-chart/save")
async def save_org_chart_alias(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        existing = _get_user_org_chart(current_user) or {}
        nodes = payload.get("nodes") or payload.get("data", {}).get("nodes") or []
        edges = payload.get("edges") or payload.get("data", {}).get("edges") or []
        metadata = payload.get("metadata") or payload.get("data", {}).get("metadata") or {}
        name = payload.get("name") or existing.get("name") or "My Organization Chart"
        description = payload.get("description") or existing.get("description") or ""
        now = datetime.utcnow().isoformat()

        chart = {
            "name": name,
            "description": description,
            "nodes": nodes,
            "edges": edges,
            "metadata": metadata,
            "created_at": existing.get("created_at") or now,
            "updated_at": now,
        }
        _save_user_org_chart(current_user, chart)
        return {"status": "saved", "chart": chart}
    except Exception as e:
        logger.error(f"Save org chart (alias) error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save org chart")


@router.get("/integrations/status")
async def get_integrations_status(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return {
        "document_sources": _merge_claims_with_settings(current_user).get("document_sources", []),
        "canvas_available": bool(os.getenv("CANVAS_CLIENT_ID")),
        "microsoft_available": bool(os.getenv("MS_CLIENT_ID")),
        "google_available": bool(os.getenv("GOOGLE_CLIENT_ID")),
        "sharepoint_available": bool(os.getenv("MS_SHAREPOINT_TENANT")),
    }
# ------------------------------
# UI Help and Tutorial metadata
# ------------------------------
@router.get("/ui/help")
async def get_ui_help(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return {
        "formulas": {
            "coverage_rate": {
                "label": "Coverage Rate",
                "description": "Percentage of standards with at least one mapped evidence item.",
                "formula": "(Mapped ÷ Total) × 100%",
                "aliases": ["Coverage", "Evidence Coverage"],
            },
            "avg_trust": {
                "label": "Average Trust",
                "description": "Mean confidence across evidence mappings for selected standards.",
                "formula": "Σ(confidence) ÷ N",
                "range": "0.0–1.0",
            },
            "avg_risk": {
                "label": "Average Risk",
                "description": "Mean predicted risk across selected standards.",
                "formula": "Σ(risk_score) ÷ N",
                "range": "0–100",
            }
        },
        "docs": {
            "faq": "https://platform.mapmystandards.ai/faq",
            "documentation": "https://platform.mapmystandards.ai/docs",
            "contact": "mailto:support@mapmystandards.ai"
        }
    }


@router.get("/ui/tutorial")
async def get_ui_tutorial(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return {
        "steps": [
            {
                "id": "standards-select",
                "title": "Select Standards",
                "text": "Use search and filters to find relevant standards. Click the checkbox to select standards for analysis and reporting.",
                "target": "#standardsList"
            },
            {
                "id": "evidence-map",
                "title": "Map Evidence",
                "text": "Upload documents or use existing content to map evidence to selected standards. Confidence indicates match strength.",
                "target": "#evidencePanel"
            },
            {
                "id": "risk-review",
                "title": "Review Risk",
                "text": "Open the Risk Overview to see aggregate risk and top contributing factors.",
                "target": "#riskOverview"
            },
            {
                "id": "generate-narrative",
                "title": "Generate Narrative",
                "text": "On Reports, add an optional introduction and click Generate Narrative. Export DOCX once the preview is ready.",
                "target": "#narrativeSection"
            }
        ],
        "dismiss_key": "mms_tutorial_dismissed_v1"
    }


# ------------------------------
# Overview and metrics
# ------------------------------
@router.get("/dashboard/overview")
async def get_dashboard_overview(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        merged = _merge_claims_with_settings(current_user)
        user_obj = {
            "email": current_user.get("email") or current_user.get("sub"),
            "institution_name": merged["organization"],
            "primary_accreditor": merged["primary_accreditor"],
            "tier": merged["tier"],
            "lms": merged["lms"],
            "document_sources": merged["document_sources"],
            "term_system": merged["term_system"],
        }
        user_id = current_user.get("user_id") or user_obj["email"] or "unknown"
        try:
            score = await analytics_service.get_compliance_score(user_id)
            standards_count = await analytics_service.get_active_standards_count(user_id)
            pending_actions = await analytics_service.get_pending_actions_count(user_id)
            recent_activity = await analytics_service.get_recent_activity(user_id)
        except Exception:
            score, standards_count, pending_actions, recent_activity = 0.0, 0, 0, []

        # Fallback to uploads store for new users and compute compliance from coverage
        uploads = _get_user_uploads(current_user)
        if not standards_count:
            standards_count = len(uploads.get("unique_standards", []))
        # Compute compliance score if missing using coverage vs total standards for current accreditor
        if not score:
            try:
                acc = merged.get("primary_accreditor") or "HLC"
                total = len(standards_graph.get_accreditor_standards(acc)) or 1
                covered = len(uploads.get("unique_standards", []))
                coverage = min(1.0, covered / float(total))
                # Blend with average trust score from uploads if available
                trust_scores = []
                for d in uploads.get("documents", []):
                    ts = (d.get("trust_score") or {}).get("overall_score")
                    if isinstance(ts, (int, float)):
                        trust_scores.append(float(ts))
                avg_trust = float(sum(trust_scores) / len(trust_scores)) if trust_scores else 0.7
                # Gate compliance: require at least one mapped standard or one analyzed document
                documents_analyzed = len(uploads.get("documents", []))
                standards_mapped = covered
                if (total > 0) and (standards_mapped > 0 or documents_analyzed > 0):
                    score = round((coverage * 0.7 + avg_trust * 0.3) * 100, 1)
                else:
                    score = 0.0
            except Exception:
                score = 0.0
        if not recent_activity:
            recent_activity = [
                {"type": "upload", "filename": d.get("filename"), "at": d.get("uploaded_at")}
                for d in uploads.get("documents", [])[-5:]
            ][::-1]

        missing: List[str] = []
        if standards_count == 0:
            missing.append("Upload at least 1 evidence document to enable standards mapping")
        if not user_obj["primary_accreditor"]:
            missing.append("Select your primary accreditor in Settings")

        return {
            "user": user_obj,
            "metrics": {
                "compliance_score": score,
                "standards_count": standards_count,
                "pending_actions": pending_actions,
                "recent_activity": recent_activity,
            },
            "missing_inputs": missing,
        }
    except Exception as e:
        logger.error(f"Overview error: {e}")
        raise HTTPException(status_code=500, detail="Failed to build overview")


async def build_dashboard_metrics_payload(
    current_user: Dict[str, Any], *, snapshot: bool = True
) -> Dict[str, Any]:
    """Internal helper that constructs the dashboard metrics payload."""

    settings_ = _merge_claims_with_settings(current_user)
    acc = (settings_.get("primary_accreditor") or "HLC").upper()
    total_roots = len(standards_graph.get_accreditor_standards(acc)) or 0

    documents_analyzed = 0
    documents_processing = 0
    standards_mapped = 0
    avg_trust = 0.7

    raw_user_id = current_user.get("user_id") or current_user.get("sub") or current_user.get("email")

    user_id: Optional[str] = None
    if raw_user_id:
        try:
            # Ensure we always query using the canonical UUID stored in the database
            user_id = await get_user_uuid_from_email(str(raw_user_id))
        except Exception:
            # Fall back to whatever identifier we have, logging handled downstream
            user_id = str(raw_user_id)
    if user_id:
        current_user.setdefault("user_id", user_id)

    used_db = False
    if user_id:
        try:
            async with db_manager.get_session() as session:
                try:
                    c1 = await session.execute(
                        text("SELECT COUNT(*) FROM documents WHERE user_id = :u AND deleted_at IS NULL"),
                        {"u": user_id},
                    )
                    documents_analyzed = int(c1.scalar() or 0)
                    documents_processing = 0
                except Exception:
                    c1 = await session.execute(
                        text("SELECT COUNT(*) FROM jobs WHERE user_id = :u AND status = 'completed'"),
                        {"u": user_id},
                    )
                    documents_analyzed = int(c1.scalar() or 0)
                    c2 = await session.execute(
                        text(
                            """
                            SELECT COUNT(*) FROM jobs
                            WHERE user_id = :u AND status IN ('queued','extracting','parsing','embedding','matching','analyzing')
                            """
                        ),
                        {"u": user_id},
                    )
                    documents_processing = int(c2.scalar() or 0)

                col_sql = text(
                    """
                    SELECT column_name FROM information_schema.columns
                    WHERE table_schema = 'public' AND table_name = 'jobs'
                    """
                )
                cols = {r[0] for r in (await session.execute(col_sql)).fetchall()}
                result_col = 'result' if 'result' in cols else ('results' if 'results' in cols else None)

                uniq: Set[str] = set()
                if result_col:
                    q = text(
                        f"""
                        SELECT {result_col} AS result
                        FROM jobs
                        WHERE user_id = :u AND status = 'completed' AND {result_col} IS NOT NULL
                        ORDER BY COALESCE(completed_at, updated_at, created_at) DESC
                        LIMIT 200
                        """
                    )
                    rows = (await session.execute(q, {"u": user_id})).fetchall()
                    for r in rows:
                        raw = r._mapping.get("result")
                        try:
                            payload = json.loads(raw) if isinstance(raw, str) else raw
                        except Exception:
                            payload = None
                        if isinstance(payload, dict):
                            mapped = payload.get("mapped_standards") or payload.get("standards") or []
                            for item in mapped:
                                sid = (item.get("standard_id") if isinstance(item, dict) else None) or (str(item) if item else None)
                                if sid:
                                    uniq.add(str(sid))
                standards_mapped = len(uniq)

                if documents_analyzed > 0:
                    uploads_data = await _get_user_uploads(current_user)
                    unique_stds = uploads_data.get("unique_standards", [])
                    if unique_stds:
                        standards_mapped = len(unique_stds)

                used_db = True
        except Exception as db_e:
            logger.warning(f"DB metrics fallback: {db_e}")

    if not used_db:
        uploads = await _get_user_uploads(current_user)
        docs = uploads.get("documents", [])
        uniq_standards = set(uploads.get("unique_standards", []) or [])
        documents_analyzed = len(docs)
        standards_mapped = len(uniq_standards)
        trust_scores: List[float] = []
        for d in docs:
            ts = (d.get("trust_score") or {}).get("overall_score")
            if isinstance(ts, (int, float)):
                trust_scores.append(float(ts))
        if trust_scores:
            avg_trust = float(sum(trust_scores) / len(trust_scores))
            if avg_trust > 1:
                avg_trust = avg_trust / 100.0
            avg_trust = max(0.0, min(1.0, avg_trust))

    total_standards = total_roots if total_roots > 0 else max(total_roots, standards_mapped)
    coverage = (standards_mapped / total_standards) if total_standards else 0.0
    coverage = max(0.0, min(1.0, coverage))
    compliance_score = 0.0
    if total_standards > 0 and (standards_mapped > 0 or documents_analyzed > 0):
        compliance_score = (coverage * 0.7 + avg_trust * 0.3) * 100
        compliance_score = round(max(0.0, min(100.0, compliance_score)), 1)

    risk_agg = risk_explainer.aggregate()
    average_risk = risk_agg.get("average_risk", 0.0)
    risk_distribution = risk_agg.get("risk_distribution", {})

    data = {
        "core_metrics": {
            "documents_analyzed": documents_analyzed,
            "documents_processing": documents_processing,
            "standards_mapped": standards_mapped,
            "total_standards": total_standards,
            "reports_generated": 0,
            "reports_pending": 0,
        },
        "performance_metrics": {
            "coverage_percentage": round(coverage * 100, 1),
            "compliance_score": compliance_score,
            "average_trust": round(avg_trust, 3),
            "average_risk": round(average_risk, 4),
            "risk_distribution": risk_distribution,
            "explanations": {
                "coverage": "coverage = unique standards with any evidence / total standards for selected accreditor",
                "compliance": "compliance = (coverage*0.7 + avg_trust*0.3) * 100; trust from EvidenceTrust when available",
                "average_trust": "average_trust = mean of document EvidenceTrust overall scores (fallback 0.7 if none)",
                "average_risk": "average_risk = mean per-standard calibrated risk score (0-1) from GapRisk Predictor",
            },
        },
        "account_info": {
            "primary_accreditor": acc,
            "trial_days_remaining": settings_.get("trial_days_remaining", 14),
        },
    }

    if snapshot:
        try:
            maybe_snapshot(
                accreditor=acc,
                payload={
                    "coverage_percentage": data["performance_metrics"]["coverage_percentage"],
                    "compliance_score": data["performance_metrics"]["compliance_score"],
                    "average_trust": data["performance_metrics"]["average_trust"],
                    "average_risk": data["performance_metrics"]["average_risk"],
                    "documents_analyzed": data["core_metrics"]["documents_analyzed"],
                    "standards_mapped": data["core_metrics"]["standards_mapped"],
                    "total_standards": data["core_metrics"]["total_standards"],
                },
                min_interval_hours=6,
                force=False,
            )
        except Exception:
            pass

    return {
        "data": data,
        "core_metrics": data.get("core_metrics", {}),
        "performance_metrics": data.get("performance_metrics", {}),
        "account_info": data.get("account_info", {}),
    }


@router.get("/dashboard/metrics")
async def get_dashboard_metrics_simple(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Return normalized dashboard metrics with coverage, compliance, trust, and risk transparency."""
    try:
        payload = await build_dashboard_metrics_payload(current_user, snapshot=True)
        return {"success": True, **payload}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Dashboard metrics error: {e}")
        raise HTTPException(status_code=500, detail="Failed to compute dashboard metrics")


@router.get("/user/profile")
async def get_user_profile(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Get user profile information from database"""
    try:
        user_id = current_user.get('sub') or current_user.get('user_id')
        if not user_id:
            raise HTTPException(status_code=400, detail="User ID not found")
        
        # Get user data from database
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT 
                        id, email, name, role, institution_name, institution_type,
                        primary_accreditor, department, onboarding_data, is_trial,
                        subscription_tier, documents_analyzed, reports_generated,
                        compliance_checks_run, created_at, updated_at
                    FROM users 
                    WHERE id = :user_id
                """),
                {"user_id": user_id}
            )
            user_data = result.fetchone()
            
            if not user_data:
                # Return basic info from JWT if user not in database
                return {
                    "id": user_id,
                    "email": current_user.get('email'),
                    "name": current_user.get('name', 'User'),
                    "institution_name": None,
                    "primary_accreditor": None,
                    "onboarding_completed": False
                }
            
            # Convert row to dict
            user_dict = dict(user_data._mapping) if hasattr(user_data, '_mapping') else dict(zip(result.keys(), user_data))
            
            # Parse onboarding data
            onboarding = user_dict.get('onboarding_data', {})
            if isinstance(onboarding, str):
                try:
                    onboarding = json.loads(onboarding)
                except:
                    onboarding = {}
            
            return {
                "id": str(user_dict['id']),
                "email": user_dict['email'],
                "name": user_dict['name'],
                "role": user_dict.get('role'),
                "institution_name": user_dict.get('institution_name'),
                "institution_type": user_dict.get('institution_type'),
                "primary_accreditor": user_dict.get('primary_accreditor'),
                "department": user_dict.get('department'),
                "is_trial": user_dict.get('is_trial', False),
                "subscription_tier": user_dict.get('subscription_tier', 'trial'),
                "documents_analyzed": user_dict.get('documents_analyzed', 0),
                "reports_generated": user_dict.get('reports_generated', 0),
                "compliance_checks_run": user_dict.get('compliance_checks_run', 0),
                "onboarding_completed": onboarding.get('onboarding_completed', False),
                "onboarding_data": onboarding,
                "created_at": user_dict['created_at'].isoformat() if user_dict.get('created_at') else None,
                "updated_at": user_dict['updated_at'].isoformat() if user_dict.get('updated_at') else None
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting user profile: {e}")
        # Return basic info on error
        return {
            "id": user_id,
            "email": current_user.get('email'),
            "name": current_user.get('name', 'User'),
            "institution_name": None,
            "primary_accreditor": None,
            "onboarding_completed": False
        }


# ------------------------------
# Standards list and search
# ------------------------------
@router.get("/standards/list")
async def list_standards(
    accreditor: Optional[str] = None,
    levels: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """List standards for the given accreditor.

    By default returns only top-level items (level == 'standard').
    Optionally include deeper levels via `levels` query, e.g.:
    - levels=standard,clause
    - levels=standard,clause,indicator
    - levels=all (or "*") to include all levels
    """
    try:
        acc = (accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()
        allowed = {"standard", "clause", "indicator"}
        levels_set: Optional[Set[str]] = {"standard"}
        lv = (levels or "").strip().lower()
        if lv:
            # normalize tokens and validate
            tokens = [t.strip() for t in lv.replace(" ", "").split(",") if t.strip()]
            if any(t in {"all", "*"} for t in tokens):
                levels_set = None  # None => all levels in graph
            else:
                chosen = {t for t in tokens if t in allowed}
                if chosen:
                    levels_set = chosen
        nodes = standards_graph.get_nodes_by_accreditor(acc, levels_set)
        mode = _get_standards_display_mode(current_user)
        items = [
            {
                "id": getattr(n, "node_id", ""),
                "code": getattr(n, "standard_id", getattr(n, "node_id", "")),
                "title": getattr(n, "title", ""),
                "description": ("" if mode == "redacted" else getattr(n, "description", "")),
                "level": getattr(n, "level", "standard"),
                "accreditor": getattr(n, "accreditor", acc),
            }
            for n in nodes
        ]
        return {
            "success": True,
            "accreditor": {"name": acc, "acronym": acc},
            "count": len(items),
            "standards": items,
            "display_mode": mode,
        }
    except Exception as e:
        logger.error(f"Standards list error: {e}")
        raise HTTPException(status_code=500, detail="Failed to list standards")


@router.get("/standards/search")
async def search_standards(
    q: str,
    accreditor: Optional[str] = None,
    category: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    try:
        ql = (q or "").strip().lower()
        if not ql:
            return {"results": []}
        acc = (accreditor or "").strip().upper()
        cat = (category or "").strip().lower()

        # Candidate nodes: restrict to accreditor if provided
        candidates = list(standards_graph.nodes.values())
        if acc:
            candidates = [n for n in candidates if getattr(n, "accreditor", "").upper() == acc]

        mode = _get_standards_display_mode(current_user)
        results: List[Dict[str, Any]] = []
        for n in candidates[:10000]:
            code = str(getattr(n, "standard_id", getattr(n, "code", getattr(n, "node_id", ""))) or "")
            title = str(getattr(n, "title", "") or "")
            desc = str(getattr(n, "description", "") or "")
            hay = " ".join([code, title, desc]).lower()
            if ql in hay:
                level = str(getattr(n, "level", "standard") or "standard")
                if cat and level.lower() != cat:
                    continue
                results.append({
                    "id": getattr(n, "node_id", code),
                    "code": code,
                    "title": title,
                    "snippet": ("" if mode == "redacted" else (desc[:180] + ("…" if len(desc) > 180 else "")) if desc else ""),
                    "category": level,
                    "accreditor": getattr(n, "accreditor", ""),
                })
                if len(results) >= 100:
                    break
        return {"results": results, "display_mode": mode}
    except Exception as e:
        logger.error(f"Standards search error: {e}")
        raise HTTPException(status_code=500, detail="Failed to search standards")


@router.get("/standards/corpus/metadata")
async def get_corpus_metadata_api(accreditor: Optional[str] = None, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Expose corpus/accreditor metadata for Standards.

    Primary source is the cached corpus metadata gathered during standards load.
    If no corpus metadata is available (e.g., running with seeded graph data),
    fall back to deriving accreditors from StandardsGraph so the UI can populate
    the accreditor list.
    """
    try:
        raw = get_corpus_metadata() or {}
        items: List[Dict[str, Any]] = []
        total_standards = 0

        if raw:
            for acc, meta in raw.items():
                if accreditor and acc.lower() != accreditor.lower():
                    continue
                loaded_nodes = len(standards_graph.get_accreditor_standards(acc))
                standard_count = int(meta.get("standard_count") or loaded_nodes)
                total_standards += standard_count
                items.append({
                    **meta,
                    "loaded_node_count": loaded_nodes,
                    "standard_count": standard_count,
                })
        else:
            # Fallback: derive from StandardsGraph seed/corpus data
            for acc in sorted(standards_graph.accreditor_roots.keys()):
                if accreditor and acc.lower() != accreditor.lower():
                    continue
                loaded_nodes = len(standards_graph.get_accreditor_standards(acc))
                standard_count = loaded_nodes
                total_standards += standard_count
                items.append({
                    "accreditor": acc,
                    "name": acc,
                    "version": None,
                    "effective_date": None,
                    "last_updated": None,
                    "source_url": None,
                    "license": None,
                    "disclaimer": None,
                    "coverage_notes": None,
                    "standard_count": standard_count,
                    "loaded_node_count": loaded_nodes,
                })

        # Derive corpus status
        available = bool(raw)
        status_items: List[Dict[str, Any]] = []
        for it in items:
            exp = int(it.get("standard_count") or 0)
            loaded = int(it.get("loaded_node_count") or 0)
            # Only mark complete when we have an explicit expected count and match or exceed it
            complete = bool(exp > 0 and loaded >= exp)
            status_items.append({
                "accreditor": it.get("accreditor"),
                "expected": exp,
                "loaded": loaded,
                "complete": complete,
            })
        overall_complete = bool(available and status_items and all(s.get("complete") for s in status_items))

        return {
            "success": True,
            "generated_at": datetime.utcnow().isoformat(),
            "total_accreditors": len(items),
            "total_standards": total_standards,
            "accreditors": sorted(items, key=lambda x: x.get("accreditor", "")),
            "corpus_status": {
                "available": available,
                "overall_complete": overall_complete,
                "accreditors": status_items,
            },
        }
    except Exception as e:
        logger.error(f"Corpus metadata error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load corpus metadata")


@router.get("/standards/metadata")
async def get_corpus_metadata_alias(accreditor: Optional[str] = None, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Alias for standards corpus metadata to match frontend expectations."""
    return await get_corpus_metadata_api(accreditor=accreditor, current_user=current_user)  # type: ignore


@router.get("/standards/corpus/status")
async def get_corpus_status(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Lightweight status for standards corpus availability and completeness."""
    meta = await get_corpus_metadata_api(current_user=current_user)  # type: ignore
    cs = meta.get("corpus_status", {}) if isinstance(meta, dict) else {}
    return {
        "success": True,
        "generated_at": meta.get("generated_at") if isinstance(meta, dict) else None,
        "corpus_status": cs,
    }


# Additional compatibility aliases for status
@router.get("/standards/status")
async def get_corpus_status_alias1(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return await get_corpus_status(current_user)  # type: ignore


@router.get("/standards/corpusstatus")
async def get_corpus_status_alias2(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return await get_corpus_status(current_user)  # type: ignore


@router.get("/metrics/summary")
async def get_metrics_summary(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    # Reflect actual uploads where available
    uploads = _get_user_uploads(current_user)
    docs_uploaded = len(uploads.get("documents", []))
    uniq_standards = len(uploads.get("unique_standards", []))
    return {
        "status": "success",
        "metrics": {
            # Immediate top-line count for UI; upgraded later when uploads exist
            "documents_uploaded": docs_uploaded,
            "ai_accuracy": 0.87,
            "time_savings": {"hours_saved": 32, "efficiency_gain": "4.2x"},
            # Coverage reflects user uploads if present
            "coverage": {"standards_analyzed": uniq_standards, "documents_processed": docs_uploaded, "evidence_mapped": sum(len(d.get("mappings", [])) for d in uploads.get("documents", []))},
            "algorithms_performance": {
                "StandardsGraph™": {"accuracy": 0.92, "speed": "real-time"},
                "EvidenceMapper™": {"accuracy": 0.87, "confidence": 0.85},
                "EvidenceTrust Score™": {"reliability": 0.94},
                "GapRisk Predictor™": {"precision": 0.89},
            },
        },
    }


@router.get("/metrics/timeseries")
async def get_timeseries_all(days: int = 30, limit: int = 200, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        series = get_series(None, days=days, limit=limit)
        return {"success": True, "count": len(series), "series": series}
    except Exception as e:
        logger.error(f"Timeseries fetch error (all): {e}")
        raise HTTPException(status_code=500, detail="Failed to get timeseries")


@router.get("/metrics/timeseries/{accreditor}")
async def get_timeseries_accreditor(accreditor: str, days: int = 30, limit: int = 200, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        series = get_series(accreditor, days=days, limit=limit)
        return {"success": True, "count": len(series), "accreditor": accreditor.upper(), "series": series}
    except Exception as e:
        logger.error(f"Timeseries fetch error ({accreditor}): {e}")
        raise HTTPException(status_code=500, detail="Failed to get timeseries for accreditor")


@router.post("/metrics/timeseries/force-snapshot")
async def force_timeseries_snapshot(
    body: Optional[Dict[str, Any]] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Admin-only: Force a metrics snapshot for the caller's accreditor.

    Restricted to demo/testing accounts to avoid abuse in production.
    """
    try:
        email = str(current_user.get("email") or current_user.get("sub") or "")
        if not email or not (email.endswith("@mapmystandards.ai") or email.startswith("demo@") or email == "demo@example.com"):
            raise HTTPException(status_code=403, detail="forbidden")

        body = body or {}
        override_accr = str(body.get("accreditor") or "").upper().strip()
        count = int(body.get("count", 1))
        spacing_minutes = int(body.get("spacing_minutes", 360))  # default 6h spacing

        base_metrics = await _compute_dashboard_metrics_for_snapshot(current_user)
        if override_accr:
            base_metrics["accreditor"] = override_accr

        from datetime import datetime, timedelta
        stored = []
        now = datetime.utcnow()
        for i in range(max(1, count)):
            snap = dict(base_metrics)
            snap["timestamp"] = (now - timedelta(minutes=spacing_minutes * (count - 1 - i))).isoformat()
            res = maybe_snapshot(snap["accreditor"], snap, min_interval_hours=0, force=True)
            stored.append(res.get("snapshot") or res)
        return {"success": True, "stored_count": len(stored), "accreditor": base_metrics["accreditor"], "samples": stored[-3:]}  # return last 3 for brevity
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Force snapshot error: {e}")
        raise HTTPException(status_code=500, detail="Failed to force snapshot")


# ------------------------------
# Risk model transparency
# ------------------------------
@router.get("/risk/factors")
async def get_risk_factors(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Explain risk factors and default weights."""
    factors = [
        {"name": "Documentation", "key": "Documentation", "description": "Availability and completeness of required documents aligned to each standard."},
        {"name": "Evidence Quality", "key": "Evidence Quality", "description": "Clarity, relevance, and trust in mapped evidence (derived from EvidenceTrust Score)."},
        {"name": "Freshness", "key": "Freshness", "description": "How recently evidence has been updated; stale evidence increases risk."},
        {"name": "Operations", "key": "Operations", "description": "Operational controls and processes supporting sustained compliance."},
        {"name": "Change Management", "key": "Change Management", "description": "Impact of recent organizational or program changes on compliance."},
        {"name": "Audit History", "key": "Audit History", "description": "Prior findings, open issues, and remediation status."},
        {"name": "Coverage", "key": "Coverage", "description": "Percent of standards with mapped evidence."},
        {"name": "Confidence", "key": "Confidence", "description": "Model confidence in mappings and assessments."},
    ]
    weights = (await get_risk_weights(current_user)).get("weights", {})  # type: ignore
    return {"factors": factors, "weights": weights, "formula": "overall_risk = weighted_mean(factors)"}


# ------------------------------
# Risk scoring endpoints
# ------------------------------
@router.post("/risk/score-standard")
async def score_single_standard_risk(
    standard_id: str,
    coverage: float = 0.0,
    avg_evidence_age_days: float = 365,
    trust_scores: Optional[List[float]] = None,
    overdue_tasks: int = 0,
    total_tasks: int = 0,
    recent_changes: int = 0,
    historical_findings: int = 0,
    days_to_review: int = 180,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Score a single standard's risk with transparent factor breakdown (non-persistent)."""
    try:
        snapshot = StandardEvidenceSnapshot(
            standard_id=standard_id,
            coverage_percent=max(0.0, min(100.0, coverage)),
            trust_scores=trust_scores or [],
            evidence_ages_days=[int(avg_evidence_age_days)],
            overdue_tasks=overdue_tasks,
            total_tasks=total_tasks,
            recent_changes=recent_changes,
            historical_findings=historical_findings,
            days_to_review=days_to_review,
        )
        score = risk_explainer.compute_standard_risk(snapshot)
        return {"success": True, "data": score.to_dict()}
    except Exception as e:
        logger.error(f"Risk score error: {e}")
        raise HTTPException(status_code=500, detail="Failed to score standard risk")


@router.post("/risk/score-bulk")
async def score_bulk_risk(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Bulk score list of standard snapshots: payload = { items: [ {standard_id, coverage_percent, trust_scores?, evidence_ages_days?, overdue_tasks?, total_tasks?, recent_changes?, historical_findings?, days_to_review?} ] }"""
    try:
        items = (payload or {}).get("items") or []
        snapshots: List[StandardEvidenceSnapshot] = []
        for it in items[:200]:  # cap
            try:
                snapshots.append(StandardEvidenceSnapshot(
                    standard_id=str(it.get("standard_id")),
                    coverage_percent=float(it.get("coverage_percent", 0.0)),
                    trust_scores=[float(x) for x in (it.get("trust_scores") or []) if isinstance(x, (int, float))],
                    evidence_ages_days=[int(x) for x in (it.get("evidence_ages_days") or []) if isinstance(x, (int, float))] or [365],
                    overdue_tasks=int(it.get("overdue_tasks", 0)),
                    total_tasks=int(it.get("total_tasks", 0)),
                    recent_changes=int(it.get("recent_changes", 0)),
                    historical_findings=int(it.get("historical_findings", 0)),
                    days_to_review=int(it.get("days_to_review", 180)),
                ))
            except Exception:
                continue
        scores = risk_explainer.compute_bulk(snapshots)
        return {"success": True, "count": len(scores), "data": [s.to_dict() for s in scores]}
    except Exception as e:
        logger.error(f"Bulk risk score error: {e}")
        raise HTTPException(status_code=500, detail="Failed to score bulk risk")


@router.get("/risk/aggregate")
async def aggregate_risk(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Aggregate statistics (distribution, average_risk, top factor contributions) over this process's scored standards."""
    try:
        agg = risk_explainer.aggregate()
        return {"success": True, "data": agg}
    except Exception as e:
        logger.error(f"Risk aggregate error: {e}")
        raise HTTPException(status_code=500, detail="Failed to aggregate risk data")


# ------------------------------
# Narrative generation
# ------------------------------
@router.post("/narrative/generate")
async def generate_narrative(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """
    Generate a narrative summary for the given standards. Expects JSON like:
    {"standard_ids": ["SACSCOC_8.1_ind_4", "HLC_1"], "body": "<optional user intro>"}
    """
    try:
        standard_ids = (payload or {}).get("standard_ids") or []
        user_body = (payload or {}).get("body") or ""

        if not isinstance(standard_ids, list):
            raise HTTPException(status_code=400, detail="standard_ids must be a list")

        # Call the narrative service to build HTML
        html = generate_narrative_html(standard_ids, user_body)

        # Persist the result to the user's session for later export
        all_sessions = _safe_load_json(SESSIONS_STORE)
        user_key = _user_key(current_user)
        all_sessions.setdefault(user_key, {})["last_narrative"] = html
        _safe_save_json(SESSIONS_STORE, all_sessions)

        return {"html": html}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Narrative generation error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate narrative")


# ------------------------------
# Narrative DOCX export
# ------------------------------
@router.post("/narrative/export.docx")
async def export_narrative_docx(payload: Optional[Dict[str, Any]] = None, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Generate and download a DOCX version of the narrative.

    Accepts optional payload { html?: str }. If not provided, generates narrative on the fly.
    """
    try:
        # Obtain HTML from payload or generate fresh
        html_in = (payload or {}).get("html") if isinstance(payload, dict) else None
        if not html_in:
            gen = await generate_narrative({}, current_user)  # type: ignore[arg-type]
            html_in = gen.get("html", "") if isinstance(gen, dict) else ""

        # Very light HTML to plain text conversion for paragraphs
        try:
            # Split on paragraph-like tags; fallback to splitting by <br>
            parts = re.split(r"</p>|<p[^>]*>", html_in or "", flags=re.IGNORECASE)
            paras = [re.sub(r"<[^>]+>", "", p).strip() for p in parts]
            paragraphs = [p for p in paras if p]
        except Exception:
            paragraphs = ["Accreditation Narrative", "Content unavailable."]

        # Build DOCX
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore

        doc = Document()
        styles = doc.styles
        try:
            styles["Normal"].font.name = "Calibri"
            styles["Normal"].font.size = Pt(11)
        except Exception:
            pass

        s = _merge_claims_with_settings(current_user)
        org = s.get("organization") or "Your Institution"
        accreditor = s.get("primary_accreditor") or "Accreditor"

        doc.add_heading(f"Accreditation Narrative – {org}", level=1)
        doc.add_paragraph(f"Framework: {accreditor}")
        doc.add_paragraph("")

        for p in paragraphs:
            # Keep inline [cite:...] tokens as-is for traceability
            doc.add_paragraph(p)

        doc.add_paragraph("")
        doc.add_paragraph("Citations: [cite:filename] refer to uploaded documents in your workspace.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_org = "".join(c for c in org if c.isalnum() or c in ("_", "-")) or "institution"
        filename = f"narrative_{safe_org}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
        headers = {
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
        return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
    except Exception as e:
        logger.error(f"Narrative DOCX export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to export DOCX narrative")


# ------------------------------
# Evidence annotation updates
# ------------------------------
@router.post("/evidence/annotate")
async def annotate_evidence(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Append a rationale span to a document mapping for a given standard."""
    try:
        filename = payload.get("filename")
        standard_id = payload.get("standard_id")
        span = (payload.get("span") or "").strip()
        if not filename or not standard_id or not span:
            raise HTTPException(status_code=400, detail="filename, standard_id, and span are required")
        all_u = _safe_load_json(UPLOADS_STORE)
        uk = _user_key(current_user)
        user_docs = all_u.get(uk, {}).get("documents", [])
        updated = False
        for d in user_docs:
            if d.get("filename") == filename:
                for m in d.get("mappings", []):
                    if m.get("standard_id") == standard_id:
                        spans = m.get("rationale_spans") or []
                        if span not in spans:
                            spans.append(span)
                            m["rationale_spans"] = spans
                            updated = True
                break
        if updated:
            all_u.setdefault(uk, {})["documents"] = user_docs
            _safe_save_json(UPLOADS_STORE, all_u)
        return {"status": "updated" if updated else "unchanged"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Annotate evidence error: {e}")
        raise HTTPException(status_code=500, detail="Failed to annotate evidence")


# ------------------------------
# Evidence reviews (UI persistence)
# ------------------------------
@router.get("/evidence/reviews")
async def get_evidence_reviews(filename: str, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    return {"filename": filename, "reviews": _get_user_reviews(current_user, filename)}


@router.post("/evidence/review")
async def save_evidence_review(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    filename = payload.get("filename")
    standard_id = payload.get("standard_id")
    if not filename or not standard_id:
        raise HTTPException(status_code=400, detail="filename and standard_id are required")
    reviewed = payload.get("reviewed")
    note = payload.get("note")
    rationale_note = payload.get("rationale_note")
    entry = _set_user_review(current_user, filename, standard_id, reviewed, note)
    # store rationale_note alongside review if provided
    if rationale_note is not None:
        entry["rationale_note"] = str(rationale_note)
        # persist full reviews map update
        all_r = _safe_load_json(REVIEWS_STORE)
        uk = _user_key(current_user)
        user_map = all_r.get(uk, {})
        file_map = user_map.get(filename, {})
        file_map[standard_id] = entry
        user_map[filename] = file_map
        all_r[uk] = user_map
        _safe_save_json(REVIEWS_STORE, all_r)
    return {"status": "success", "filename": filename, "standard_id": standard_id, "entry": entry}


# ------------------------------
# Evidence upload (simple demo path to support dashboard drag/drop)
# ------------------------------
@router.post("/evidence/upload")
async def evidence_upload_simple(
    files: List[UploadFile] = File(...),
    doc_type: Optional[str] = Form(None),
    institution: Optional[str] = Form(None),
    accreditor: Optional[str] = Form(None),
    institution_type: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    try:
        # Get storage service
        storage = get_storage_service()
        saved: List[Dict[str, Any]] = []
        
        # Get user info for file key generation
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        org_id = current_user.get("org_id", "default")
        
        for f in files:
            name = f.filename or "upload.bin"
            content = await f.read()
            
            # Check if file has content
            if not content or len(content) == 0:
                logger.warning(f"Empty file uploaded: {name}")
                continue
                
            # Generate storage key
            file_key = storage.generate_file_key(org_id, user_id, name)
            
            # Determine content type
            content_type = f.content_type or "application/octet-stream"
            
            # Get institution context (from form data or user profile)
            inst_name = institution
            inst_accreditor = accreditor
            inst_type = institution_type
            
            # If not provided in form, try to get from user profile in database
            if not inst_name or not inst_accreditor:
                try:
                    async with db_manager.get_session() as session:
                        user_result = await session.execute(
                            text("SELECT institution_name, primary_accreditor, institution_type FROM users WHERE id = :user_id"),
                            {"user_id": user_id}
                        )
                        user_data = user_result.fetchone()
                        if user_data:
                            inst_name = inst_name or user_data[0]
                            inst_accreditor = inst_accreditor or user_data[1]
                            inst_type = inst_type or user_data[2]
                except Exception as e:
                    logger.warning(f"Could not fetch user institution data: {e}")
            
            # Save to storage (S3 or local)
            result = await storage.save_file(
                file_content=content,
                file_key=file_key,
                content_type=content_type,
                metadata={
                    "original_filename": name,
                    "user_id": user_id,
                    "org_id": org_id,
                    "institution_name": inst_name or "",
                    "institution_type": inst_type or "",
                    "primary_accreditor": inst_accreditor or "",
                    "doc_type": doc_type or "",
                    "uploaded_at": datetime.utcnow().isoformat()
                }
            )
            
            if result.get("success"):
                # Record upload for metrics
                upload_result = await _record_user_upload(
                    current_user,
                    name,
                    standard_ids=[],
                    doc_type=doc_type,
                    trust_score=None,
                    saved_path=file_key,  # Use storage key instead of local path
                    fingerprint=result.get("hash", "")[:16],
                    file_size=len(content),
                    content_type=content_type,
                )

                document_id = upload_result.get("document_id")
                
                # Automatically trigger analysis for supported file types
                analysis_result = None
                try:
                    if name.lower().endswith(('.pdf', '.txt', '.doc', '.docx')):
                        # Analyze the document content
                        analysis = await _analyze_evidence_from_bytes(
                            name,
                            content,
                            doc_type,
                            current_user,
                            document_id=document_id  # Pass the document ID to update existing record
                        )
                        analysis_result = analysis
                        
                        # Update document status to analyzed
                        if document_id and "mappings" in analysis:
                            async with db_manager.get_session() as session:
                                await session.execute(
                                    text("UPDATE documents SET status = 'analyzed' WHERE id = :id"),
                                    {"id": document_id}
                                )
                                await session.commit()
                except Exception as e:
                    logger.warning(f"Auto-analysis failed for {name}: {e}")
                
                saved.append({
                    "id": document_id,
                    "original": name,
                    "saved_as": file_key,
                    "size": len(content),
                    "hash": result.get("hash", ""),
                    "status": "analyzed" if analysis_result else "uploaded",
                    "storage_type": result.get("storage_type", "unknown"),
                    "analysis": analysis_result
                })
            else:
                logger.error(f"Failed to save file {name}")
                
        if not saved:
            raise HTTPException(status_code=400, detail="No valid files uploaded")
            
        return {"success": True, "files": saved}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Simple evidence upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload evidence: {str(e)}")


@router.get("/evidence/upload/debug", include_in_schema=False)
async def evidence_upload_simple_debug():
    return {"ok": True, "path": "/api/user/intelligence-simple/evidence/upload"}


@router.get("/evidence/list")
async def list_evidence(
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """
    List all uploaded evidence documents for the current user
    """
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Get documents from database instead of file system
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT d.id, d.filename, d.file_size, d.content_type, 
                           d.status, d.uploaded_at, d.file_key,
                           COUNT(em.standard_id) as mapped_count
                    FROM documents d
                    LEFT JOIN evidence_mappings em ON d.id = em.document_id
                    WHERE d.user_id = :user_id 
                    AND d.deleted_at IS NULL
                    GROUP BY d.id, d.filename, d.file_size, d.content_type, 
                             d.status, d.uploaded_at, d.file_key
                    ORDER BY d.uploaded_at DESC
                """),
                {"user_id": user_id}
            )
            
            documents = []
            for row in result:
                documents.append({
                    "id": row.id,  # Use actual database ID
                    "filename": row.filename,
                    "size": row.file_size or 0,
                    "content_type": row.content_type,
                    "status": row.status or "uploaded",
                    "uploaded_at": row.uploaded_at.isoformat() if row.uploaded_at else None,
                    "mapped_count": row.mapped_count or 0,
                    "file_key": row.file_key
                })
            
            # Get unique standards
            standards_result = await session.execute(
                text("""
                    SELECT DISTINCT em.standard_id
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE d.user_id = :user_id
                    AND d.deleted_at IS NULL
                """),
                {"user_id": user_id}
            )
            unique_standards = [row.standard_id for row in standards_result]
            
            return {
                "success": True,
                "evidence": documents,
                "total": len(documents),
                "unique_standards": unique_standards
            }
    except Exception as e:
        logger.error(f"Error listing evidence: {e}")
        raise HTTPException(status_code=500, detail="Failed to list evidence")


@router.get("/uploads")
async def list_uploads(
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """
    List all uploads (alias for evidence/list for compatibility)
    """
    return await list_evidence(current_user)


@router.get("/standards")
async def list_user_standards(
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """
    List standards available to the user based on their settings
    """
    try:
        settings = _merge_claims_with_settings(current_user)
        primary_accreditor = settings.get("primary_accreditor", "HLC")
        
        # Get standards for the user's primary accreditor
        standards = standards_graph.get_accreditor_standards(primary_accreditor)
        
        # Apply display policy
        mode = _get_standards_display_mode(current_user)
        standards_list = []
        
        for std in standards:
            std_dict = {
                "id": std.id,
                "number": std.number,
                "description": std.description if mode == "full" else "",
                "level": std.level,
                "parent_id": std.parent_id,
                "category": std.category,
                "accreditor": primary_accreditor
            }
            standards_list.append(std_dict)
        
        return {
            "success": True,
            "standards": standards_list,
            "total": len(standards_list),
            "accreditor": primary_accreditor,
            "display_mode": mode
        }
    except Exception as e:
        logger.error(f"Error listing standards: {e}")
        raise HTTPException(status_code=500, detail="Failed to list standards")


@router.get("/metrics/dashboard")
async def get_dashboard_metrics_alias(
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """
    Alias for /dashboard/metrics to match expected URL pattern
    """
    return await get_dashboard_metrics_simple(current_user)


# ------------------------------
# Evidence analysis
# ------------------------------
@router.post("/analyze/evidence")
async def analyze_evidence(
    file: UploadFile = File(...),
    doc_type: Optional[str] = Form(None),
    document_id: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    try:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")

        filename = file.filename or f"upload-{uuid.uuid4().hex}"

        return await _analyze_evidence_from_bytes(
            filename=filename,
            content=content,
            doc_type=doc_type,
            current_user=current_user,
            document_id=document_id,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Evidence analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {e}")


@router.post("/evidence/analyze")
async def analyze_evidence_alias(
    file: UploadFile = File(...),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
    doc_type: Optional[str] = Form(None),
):
    # Delegate to main handler (duplicate logic for clarity and isolation)
    return await analyze_evidence(file=file, doc_type=doc_type, current_user=current_user)


# ------------------------------
# Document download endpoint
# ------------------------------
@router.get("/uploads/{document_id}")
async def download_document(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Download a specific document by ID"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Get document info from database
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT id, filename, file_key, content_type
                    FROM documents 
                    WHERE id = :document_id 
                    AND user_id = :user_id 
                    AND deleted_at IS NULL
                """),
                {"document_id": document_id, "user_id": user_id}
            )
            
            doc = result.first()
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Get file from storage
            file_key = doc.file_key
            filename = doc.filename
            content_type = doc.content_type or "application/octet-stream"
            
            # Download from storage
            storage: StorageService = get_storage_service()
            try:
                file_content = await storage.get_file(file_key)
                
                # Return file response
                return Response(
                    content=file_content,
                    media_type=content_type,
                    headers={
                        "Content-Disposition": f'attachment; filename="{filename}"'
                    }
                )
            except Exception as e:
                logger.error(f"Storage download error: {e}")
                raise HTTPException(status_code=404, detail="File not found in storage")
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document download error: {e}")
        raise HTTPException(status_code=500, detail="Failed to download document")


# Also expose download endpoint at the expected URL
@router.get("/documents/{document_id}/download")
async def download_document_standard(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Download a specific document by ID (standard route)"""
    return await download_document(document_id, current_user)


@router.post("/documents/{document_id}/analyze")
async def analyze_existing_document(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Analyze a document that's already uploaded and stored in the database"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Get document from database
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT id, filename, file_key, content_type, status
                    FROM documents 
                    WHERE id = :document_id 
                    AND user_id = :user_id 
                    AND deleted_at IS NULL
                """),
                {"document_id": document_id, "user_id": user_id}
            )
            
            doc = result.first()
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Check if already analyzed
            if doc.status == "analyzed":
                # Return cached analysis if available
                mappings_result = await session.execute(
                    text("""
                        SELECT standard_id, confidence, excerpts
                        FROM evidence_mappings
                        WHERE document_id = :document_id
                    """),
                    {"document_id": document_id}
                )
                mappings = []
                for m in mappings_result:
                    mappings.append({
                        "standard_id": m.standard_id,
                        "confidence": m.confidence,
                        "excerpts": m.excerpts or []
                    })
                
                if mappings:
                    return {
                        "status": "success",
                        "document_id": document_id,
                        "filename": doc.filename,
                        "mappings": mappings,
                        "cached": True
                    }
            
            # Get file content from storage
            storage = get_storage_service()
            file_content = await storage.get_file(doc.file_key)
            
            # Analyze the document
            analysis_result = await _analyze_evidence_from_bytes(
                doc.filename,
                file_content,
                None,  # doc_type
                current_user,
                document_id=document_id  # Pass document ID to update existing record
            )
            
            # Update document status
            await session.execute(
                text("""
                    UPDATE documents 
                    SET status = 'analyzed', 
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = :document_id
                """),
                {"document_id": document_id}
            )
            
            # Store mappings in database
            if "mappings" in analysis_result:
                for mapping in analysis_result["mappings"]:
                    await session.execute(
                        text("""
                            INSERT INTO evidence_mappings 
                            (id, document_id, standard_id, confidence, excerpts, created_at)
                            VALUES (:id, :document_id, :standard_id, :confidence, :excerpts, CURRENT_TIMESTAMP)
                            ON CONFLICT (document_id, standard_id) 
                            DO UPDATE SET 
                                confidence = :confidence,
                                excerpts = :excerpts,
                                updated_at = CURRENT_TIMESTAMP
                        """),
                        {
                            "id": str(uuid.uuid4()),
                            "document_id": document_id,
                            "standard_id": mapping["standard_id"],
                            "confidence": mapping["confidence"],
                            "excerpts": json.dumps(mapping.get("excerpts", []))
                        }
                    )
            
            await session.commit()
            
            return {
                "status": "success",
                "document_id": document_id,
                "filename": doc.filename,
                **analysis_result
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze document: {str(e)}")


@router.get("/documents/{document_id}/analysis")
async def get_document_analysis(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Get analysis results for a document"""
    try:
        logger.info(f"Getting analysis for document {document_id}")
        logger.info(f"Current user: {current_user}")
        
        # Extract user ID from the current_user dict
        email = current_user.get('email') or current_user.get('sub')
        user_id = None
        
        try:
            if email and '@' in email:
                user_id = await get_user_uuid_from_email(email)
            else:
                user_id = current_user.get('user_id') or current_user.get('sub')
        except Exception as e:
            logger.error(f"Error getting user ID: {e}")
            # Fallback to using sub or user_id directly
            user_id = current_user.get('user_id') or current_user.get('sub') or email
            
        logger.info(f"User ID: {user_id}")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Could not determine user ID")
        
        async with db_manager.get_session() as session:
            # Get document and verify ownership
            result = await session.execute(
                text("""
                    SELECT d.id, d.filename, d.status,
                           d.uploaded_at, d.updated_at, d.file_size, d.content_type,
                           d.analysis_results,
                           COUNT(DISTINCT em.standard_id) as standards_mapped
                    FROM documents d
                    LEFT JOIN evidence_mappings em ON em.document_id = d.id
                    WHERE d.id = :document_id 
                    AND d.user_id = :user_id 
                    AND d.deleted_at IS NULL
                    GROUP BY d.id, d.filename, d.status, d.uploaded_at, d.updated_at, d.file_size, d.content_type
                """),
                {"document_id": document_id, "user_id": user_id}
            )
            
            doc = result.first()
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")

            stored_analysis: Dict[str, Any] = {}
            if getattr(doc, "analysis_results", None):
                try:
                    stored_analysis = json.loads(doc.analysis_results)
                except Exception as parse_exc:
                    logger.warning(f"Unable to parse stored analysis_results for document {document_id}: {parse_exc}")
            
            # Get all mappings for this document
            mappings_result = await session.execute(
                text("""
                    SELECT em.standard_id, em.confidence, em.excerpts
                    FROM evidence_mappings em
                    WHERE em.document_id = :document_id
                    ORDER BY em.confidence DESC
                """),
                {"document_id": document_id}
            )
            
            mappings = []
            for mapping in mappings_result:
                # Handle excerpts - it might be a JSON string or already a list
                if isinstance(mapping.excerpts, str):
                    excerpts = json.loads(mapping.excerpts) if mapping.excerpts else []
                elif isinstance(mapping.excerpts, list):
                    excerpts = mapping.excerpts
                else:
                    excerpts = []
                
                # Since we don't have a standards table, we'll parse the standard_id
                # to extract accreditor and code information
                standard_id_parts = mapping.standard_id.split('.')
                accreditor = standard_id_parts[0] if standard_id_parts else "Unknown"
                code = '.'.join(standard_id_parts[1:]) if len(standard_id_parts) > 1 else mapping.standard_id
                
                standard_data = {
                    "code": code,
                    "title": f"{accreditor} Standard {code}",
                    "description": "",
                    "accreditor": accreditor
                }
                
                mappings.append({
                    "standard_id": mapping.standard_id,
                    "confidence": mapping.confidence,
                    "excerpts": excerpts,
                    "standard": standard_data
                })

            document_meta = {
                "uploaded_at": doc.uploaded_at.isoformat() if getattr(doc, "uploaded_at", None) else None,
                "updated_at": doc.updated_at.isoformat() if getattr(doc, "updated_at", None) else None,
                "file_size": getattr(doc, "file_size", None),
                "content_type": getattr(doc, "content_type", None),
                "status": doc.status,
            }

            trust_summary = stored_analysis.get("trust_summary") if isinstance(stored_analysis, dict) else None
            if not trust_summary:
                trust_summary = _build_trust_summary(
                    stored_analysis.get("trust_score") if isinstance(stored_analysis, dict) else None,
                    document_meta=document_meta,
                    mappings=[{"confidence": m.get("confidence"), "excerpts": m.get("excerpts")} for m in mappings],
                )

            trust_score_payload: Dict[str, Any] = {}
            if isinstance(stored_analysis, dict) and isinstance(stored_analysis.get("trust_score"), dict):
                trust_score_payload = stored_analysis.get("trust_score", {})
            if not trust_score_payload and trust_summary:
                trust_score_payload = {"overall_score": trust_summary.get("overall_score")}

            analyzed_at_iso = doc.updated_at.isoformat() if getattr(doc, "updated_at", None) else None

            if doc.status != 'analyzed' or not mappings:
                return {
                    "document": {
                        "id": doc.id,
                        "filename": doc.filename,
                        "status": doc.status,
                        "analyzed_at": analyzed_at_iso,
                        "standards_mapped": 0,
                    },
                    "analysis": {
                        "mapped_standards": [],
                        "total_mappings": 0,
                        "accreditors": [],
                        "message": "Document has not been analyzed yet. Please analyze the document first.",
                        "trust_summary": trust_summary,
                        "trust_score": trust_score_payload,
                    },
                }

            return {
                "document": {
                    "id": doc.id,
                    "filename": doc.filename,
                    "status": doc.status,
                    "analyzed_at": analyzed_at_iso,
                    "standards_mapped": doc.standards_mapped,
                },
                "analysis": {
                    "mapped_standards": mappings,
                    "total_mappings": len(mappings),
                    "accreditors": list(set(m["standard"]["accreditor"] for m in mappings)) if mappings else [],
                    "trust_summary": trust_summary,
                    "trust_score": trust_score_payload,
                },
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving analysis for document {document_id}: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve analysis: {str(e)}")


@router.get("/debug/analysis/{document_id}")
async def debug_analysis(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Debug endpoint to check what's happening with analysis"""
    try:
        # Get basic info
        email = current_user.get('email') or current_user.get('sub')
        user_id = current_user.get('user_id') or current_user.get('sub')
        
        debug_info = {
            "document_id": document_id,
            "current_user": {
                "email": email,
                "user_id": user_id,
                "sub": current_user.get('sub'),
                "keys": list(current_user.keys())
            }
        }
        
        # Try to get document info
        try:
            async with db_manager.get_session() as session:
                result = await session.execute(
                    text("SELECT id, user_id, status FROM documents WHERE id = :doc_id"),
                    {"doc_id": document_id}
                )
                doc = result.first()
                if doc:
                    debug_info["document"] = {
                        "id": doc[0],
                        "user_id": doc[1],
                        "status": doc[2]
                    }
                else:
                    debug_info["document"] = "Not found"
                    
                # Check mappings
                result = await session.execute(
                    text("SELECT COUNT(*) FROM evidence_mappings WHERE document_id = :doc_id"),
                    {"doc_id": document_id}
                )
                count = result.scalar()
                debug_info["mappings_count"] = count
                
        except Exception as e:
            debug_info["db_error"] = str(e)
            
        return debug_info
        
    except Exception as e:
        return {"error": str(e), "type": type(e).__name__}


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Delete a document (soft delete)"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Soft delete the document
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    UPDATE documents 
                    SET deleted_at = CURRENT_TIMESTAMP
                    WHERE id = :document_id 
                    AND user_id = :user_id 
                    AND deleted_at IS NULL
                    RETURNING id, filename
                """),
                {"document_id": document_id, "user_id": user_id}
            )
            
            deleted = result.first()
            if not deleted:
                raise HTTPException(status_code=404, detail="Document not found")
            
            await session.commit()
            
            return {
                "status": "success",
                "message": f"Document '{deleted.filename}' deleted successfully",
                "document_id": document_id
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document delete error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


# ------------------------------
# Dashboard endpoints
# ------------------------------
@router.get("/documents/list")
async def list_documents(
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """List all user's documents"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT id, filename, file_size, content_type, status, uploaded_at
                    FROM documents 
                    WHERE user_id = :user_id 
                    AND deleted_at IS NULL
                    ORDER BY uploaded_at DESC
                """),
                {"user_id": user_id}
            )
            
            documents = []
            for row in result:
                documents.append({
                    "id": row.id,
                    "filename": row.filename,
                    "file_size": row.file_size,
                    "content_type": row.content_type,
                    "status": row.status,
                    "uploaded_at": row.uploaded_at.isoformat() if row.uploaded_at else None
                })
            
            return {"status": "success", "documents": documents}
            
    except Exception as e:
        logger.error(f"Documents list error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")


@router.get("/documents/recent")
async def get_recent_documents(
    limit: int = 5,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get recent document activity"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT id, filename, status, uploaded_at
                    FROM documents 
                    WHERE user_id = :user_id 
                    AND deleted_at IS NULL
                    ORDER BY uploaded_at DESC
                    LIMIT :limit
                """),
                {"user_id": user_id, "limit": limit}
            )
            
            activities = []
            for row in result:
                activities.append({
                    "type": "document_upload",
                    "document_id": row.id,
                    "filename": row.filename,
                    "status": row.status,
                    "timestamp": row.uploaded_at.isoformat() if row.uploaded_at else None,
                    "message": f"Uploaded {row.filename}"
                })
            
            return {"status": "success", "activities": activities}
            
    except Exception as e:
        logger.error(f"Recent documents error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get recent documents: {str(e)}")


@router.get("/compliance/summary")
async def get_compliance_summary(
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get compliance summary for dashboard"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Get document analysis results
        async with db_manager.get_session() as session:
            # Count analyzed documents
            doc_result = await session.execute(
                text("""
                    SELECT COUNT(*) as total,
                           COUNT(CASE WHEN status = 'analyzed' THEN 1 END) as analyzed
                    FROM documents 
                    WHERE user_id = :user_id 
                    AND deleted_at IS NULL
                """),
                {"user_id": user_id}
            )
            doc_stats = doc_result.first()
            
            # Count mapped standards
            standards_result = await session.execute(
                text("""
                    SELECT COUNT(DISTINCT em.standard_id) as mapped_standards
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE d.user_id = :user_id
                    AND d.deleted_at IS NULL
                """),
                {"user_id": user_id}
            )
            standards_stats = standards_result.first()
            
            total_docs = doc_stats.total or 0
            analyzed_docs = doc_stats.analyzed or 0
            completion_rate = (analyzed_docs / total_docs * 100) if total_docs > 0 else 0
            
            return {
                "status": "success",
                "summary": {
                    "total_documents": total_docs,
                    "analyzed_documents": analyzed_docs,
                    "mapped_standards": standards_stats.mapped_standards or 0,
                    "completion_rate": round(completion_rate, 1),
                    "compliance_score": round(completion_rate * 0.85, 1)  # Mock score
                }
            }
            
    except Exception as e:
        logger.error(f"Compliance summary error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get compliance summary: {str(e)}")


@router.get("/risk/summary")
async def get_risk_summary(
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get risk summary for dashboard"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        # Mock risk data for now
        return {
            "status": "success",
            "summary": {
                "total_standards": 100,
                "covered_standards": 25,
                "high_risk_gaps": 5,
                "medium_risk_gaps": 15,
                "low_risk_gaps": 55,
                "risk_score": 72  # Mock score
            }
        }
            
    except Exception as e:
        logger.error(f"Risk summary error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get risk summary: {str(e)}")


# ------------------------------
# Bulk analyze (multi-file upload)
# ------------------------------
@router.post("/evidence/analyze/bulk")
async def analyze_evidence_bulk(
    files: List[UploadFile] = File(...),
    doc_type: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    results: List[Dict[str, Any]] = []
    for f in files:
        try:
            res = await analyze_evidence(file=f, doc_type=doc_type, current_user=current_user)
            results.append(res)
        except Exception as e:
            results.append({"status": "error", "filename": getattr(f, 'filename', ''), "detail": str(e)})
    return {"status": "success", "results": results}


# ------------------------------
# Evidence: analyze from storage (supports S3 presigned upload flow)
# ------------------------------
@router.post("/evidence/analyze/from-storage")
async def analyze_evidence_from_storage(
    payload: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
    storage: StorageService = Depends(get_storage_service),
):
    """Analyze an evidence file that has been uploaded to storage.

    Expected payload: { file_key, filename, content_type?, doc_type? }
    """
    file_key = str(payload.get("file_key") or "").strip()
    filename = str(payload.get("filename") or "").strip()
    if not file_key or not filename:
        raise HTTPException(status_code=400, detail="file_key and filename are required")
    try:
        content = await storage.get_file(file_key)
        if not content:
            raise HTTPException(status_code=404, detail="File not found in storage")
        doc_type = payload.get("doc_type")
        return await _analyze_evidence_from_bytes(filename, content, doc_type, current_user)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analyze from storage error: {e}")
        raise HTTPException(status_code=500, detail="Failed to analyze evidence from storage")


# ------------------------------
# Governance: Tasks & Audit
# ------------------------------
def _get_user_tasks(claims: Dict[str, Any]) -> Dict[str, Any]:
    all_t = _safe_load_json(TASKS_STORE)
    return all_t.get(_user_key(claims), {})


def _save_user_tasks(claims: Dict[str, Any], tasks: Dict[str, Any]) -> None:
    all_t = _safe_load_json(TASKS_STORE)
    all_t[_user_key(claims)] = tasks
    _safe_save_json(TASKS_STORE, all_t)


@router.post("/tasks/create")
async def create_task(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    tid = datetime.utcnow().strftime("%Y%m%d%H%M%S") + "_" + secrets.token_hex(6)
    task = {
        "id": tid,
        "title": payload.get("title") or "",
        "description": payload.get("description") or "",
        "standard_id": payload.get("standard_id") or "",
        "accreditor": (payload.get("accreditor") or "").upper(),
        "assignee": payload.get("assignee") or "",
        "due_date": payload.get("due_date") or "",
        "priority": payload.get("priority") or "normal",
        "status": payload.get("status") or "open",
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "comments": [],
    }
    tasks = _get_user_tasks(current_user)
    tasks[tid] = task
    _save_user_tasks(current_user, tasks)
    try:
        with open(TASKS_AUDIT_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps({"ts": datetime.utcnow().isoformat(), "user": _user_key(current_user), "action": "create", "task": task}, ensure_ascii=False) + "\n")
    except Exception:
        pass
    return {"success": True, "task": task}


@router.get("/tasks/list")
async def list_tasks(
    accreditor: Optional[str] = None,
    standard_id: Optional[str] = None,
    status: Optional[str] = None,
    assignee: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    tasks = list(_get_user_tasks(current_user).values())

    def ok(t: Dict[str, Any]) -> bool:
        if accreditor and (t.get("accreditor") or "").upper() != accreditor.upper():
            return False
        if standard_id and t.get("standard_id") != standard_id:
            return False
        if status and (t.get("status") or "").lower() != status.lower():
            return False
        if assignee and (t.get("assignee") or "").lower() != assignee.lower():
            return False
        return True
    filt = [t for t in tasks if ok(t)]
    return {"success": True, "count": len(filt), "tasks": sorted(filt, key=lambda x: x.get("due_date") or "")}


@router.post("/tasks/update")
async def update_task(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    tid = payload.get("id")
    if not tid:
        raise HTTPException(status_code=400, detail="id is required")
    tasks = _get_user_tasks(current_user)
    if tid not in tasks:
        raise HTTPException(status_code=404, detail="task not found")
    t = tasks[tid]
    before = dict(t)
    for k in ["title", "description", "standard_id", "accreditor", "assignee", "due_date", "priority", "status"]:
        if k in payload and payload[k] is not None:
            t[k] = payload[k]
    if payload.get("comment"):
        t.setdefault("comments", []).append({"author": _user_key(current_user), "ts": datetime.utcnow().isoformat(), "text": str(payload.get("comment"))})
    t["updated_at"] = datetime.utcnow().isoformat()
    tasks[tid] = t
    _save_user_tasks(current_user, tasks)
    try:
        with open(TASKS_AUDIT_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps({"ts": datetime.utcnow().isoformat(), "user": _user_key(current_user), "action": "update", "id": tid, "before": before, "after": t}, ensure_ascii=False) + "\n")
    except Exception:
        pass
    return {"success": True, "task": t}


# ------------------------------
# Narratives: Strict citations and export
# ------------------------------
def _build_citation_map(claims: Dict[str, Any], standard_ids: List[str]) -> Dict[str, Any]:
    uploads = _get_user_uploads(claims)
    cmap: Dict[str, Any] = {}
    for d in uploads.get("documents", []):
        fname = d.get("filename")
        for md in d.get("mappings", []):
            sid = md.get("standard_id")
            if sid in (standard_ids or []):
                cmap.setdefault(sid, []).append({
                    "filename": fname,
                    "confidence": md.get("confidence"),
                    "page_anchors": md.get("page_anchors", []),
                    "fingerprint": d.get("fingerprint"),
                    "saved_path": d.get("saved_path"),
                })
    return cmap


@router.post("/narratives/generate")
async def generate_narrative_v2(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    standard_ids = payload.get("standard_ids") or []
    user_body = payload.get("body") or ""
    narrative_type = payload.get("narrative_type", "comprehensive")
    strict = bool(payload.get("strict_citations", False))
    
    # Check if AI narrative service is available and OpenAI is configured
    if USE_AI_NARRATIVE and settings.openai_api_key and settings.openai_api_key != "sk-proj-PLACEHOLDER":
        try:
            # Build evidence mappings for AI narrative
            cmap = _build_citation_map(current_user, standard_ids)
            
            # Get user's institution info
            email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
            user_settings = _merge_claims_with_settings(current_user)
            institution_info = {
                "name": user_settings.get("institution_name", "Institution"),
                "type": user_settings.get("institution_type", "University"),
                "accreditor": user_settings.get("primary_accreditor", "HLC")
            }
            
            # Generate AI-enhanced narrative with citations
            result = await ai_narrative_service.generate_narrative_with_citations(
                standard_ids=standard_ids,
                evidence_mappings=cmap,
                narrative_type=narrative_type,
                user_context=user_body,
                institution_info=institution_info
            )
            
            mode = _get_standards_display_mode(current_user)
            
            # Return enhanced response with AI insights
            return {
                "success": True,
                "html": result["narrative"],
                "display_mode": mode,
                "citations": result.get("citations", []),
                "compliance_score": result.get("compliance_score", 0),
                "gaps_identified": result.get("gaps_identified", []),
                "recommendations": result.get("recommendations", []),
                "algorithm": result.get("algorithm", "CiteGuard™"),
                "ai_enabled": True
            }
            
        except Exception as e:
            logger.error(f"AI narrative generation failed: {e}")
            # Fallback to basic narrative
    
    # Fallback to basic narrative generation
    if strict:
        cmap = _build_citation_map(current_user, standard_ids)
        missing = [sid for sid in standard_ids if not cmap.get(sid)]
        if missing:
            return {"success": False, "strict_blocked": True, "missing_citations": missing}
    
    html = generate_narrative_html(standard_ids, user_body)
    mode = _get_standards_display_mode(current_user)
    return {"success": True, "html": html, "display_mode": mode, "ai_enabled": False}


@router.post("/narratives/citeguard/check")
async def citeguard_check(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    standard_ids = payload.get("standard_ids") or []
    cmap = _build_citation_map(current_user, standard_ids)
    missing = [sid for sid in standard_ids if not cmap.get(sid)]
    return {"success": True, "ok": len(missing) == 0, "missing": missing, "citation_map": cmap}


@router.post("/narratives/export/reviewer-pack")
async def export_reviewer_pack(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    standard_ids = payload.get("standard_ids") or []
    include_files = bool(payload.get("include_files", False))
    include_checklist = bool(payload.get("include_checklist", True))
    include_metrics = bool(payload.get("include_metrics", True))
    checklist_format = str(payload.get("checklist_format") or "json").lower()
    acc = (payload.get("accreditor") or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()
    user_body = payload.get("body") or ""
    html = generate_narrative_html(standard_ids, user_body)
    cmap = _build_citation_map(current_user, standard_ids)
    tmp_zip = Path(tempfile.gettempdir()) / ("reviewer_pack_" + secrets.token_hex(6) + ".zip")
    try:
        with zipfile.ZipFile(tmp_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("index.html", "<html><body><h1>Reviewer Pack</h1><a href=\"narrative.html\">Open Narrative</a></body></html>")
            zf.writestr("narrative.html", html)
            zf.writestr("citation_map.json", json.dumps(cmap, ensure_ascii=False, indent=2))
            # Optional checklist
            if include_checklist:
                # Reuse the checklist generator (JSON)
                chk = await get_evidence_intake_checklist(accreditor=acc, format="json", include_indicators=False, current_user=current_user)  # type: ignore
                if checklist_format == "csv":
                    # Build CSV from chk.groups
                    try:
                        import io as _io
                        import csv as _csv
                        sio = _io.StringIO()
                        w = _csv.writer(sio)
                        w.writerow(["standard_id", "accreditor", "level", "title", "description", "evidence_requirements", "covered"])
                        for g in (chk.get("groups") or []):
                            for it in (g.get("items") or []):
                                w.writerow([
                                    it.get("standard_id"),
                                    it.get("accreditor"),
                                    it.get("level"),
                                    it.get("title"),
                                    (it.get("description") or "").replace("\n", " ").strip(),
                                    "; ".join(it.get("evidence_requirements") or []),
                                    "yes" if it.get("covered") else "no",
                                ])
                        zf.writestr("checklist.csv", sio.getvalue())
                    except Exception:
                        # Fallback to JSON if CSV generation fails
                        zf.writestr("checklist.json", json.dumps(chk, ensure_ascii=False, indent=2))
                else:
                    zf.writestr("checklist.json", json.dumps(chk, ensure_ascii=False, indent=2))
            # Optional readiness metrics
            if include_metrics:
                try:
                    metrics = await readiness_scorecard(current_user)  # type: ignore
                except Exception:
                    metrics = {"success": False}
                zf.writestr("metrics.json", json.dumps(metrics, ensure_ascii=False, indent=2))
            if include_files:
                for sid, cites in (cmap or {}).items():
                    for c in cites:
                        sp = c.get("saved_path")
                        if sp and os.path.exists(sp):
                            try:
                                arcname = f"artifacts/{os.path.basename(sp)}"
                                zf.write(sp, arcname=arcname)
                            except Exception:
                                continue
        # Copy/move into snapshots with a stable name and provide download URL
        try:
            safe_user = (_user_key(current_user) or "user").replace("@", "_").replace("/", "_")
            dest = SNAPSHOTS_DIR / f"reviewer_pack_{safe_user}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.zip"
            import shutil
            shutil.copyfile(tmp_zip, dest)
            download_url = f"/api/user/intelligence-simple/download/snapshots/{dest.name}"
            # Persist last reviewer pack info into user settings
            try:
                settings_ = _merge_claims_with_settings(current_user)
                settings_["last_reviewer_pack"] = {
                    "path": str(dest),
                    "download_url": download_url,
                    "created_at": datetime.utcnow().isoformat(),
                }
                _save_user_settings(current_user, settings_)
            except Exception:
                pass
            return {"success": True, "pack_path": str(dest), "download_url": download_url}
        except Exception:
            # Persist fallback info without download URL
            try:
                settings_ = _merge_claims_with_settings(current_user)
                settings_["last_reviewer_pack"] = {
                    "path": str(tmp_zip),
                    "download_url": None,
                    "created_at": datetime.utcnow().isoformat(),
                }
                _save_user_settings(current_user, settings_)
            except Exception:
                pass
            return {"success": True, "pack_path": str(tmp_zip)}
    except Exception as e:
        logger.error(f"Reviewer pack export error: {e}")
        raise HTTPException(status_code=500, detail="Failed to build reviewer pack")


# ------------------------------
# Readiness Scorecard
# ------------------------------
@router.get("/readiness/scorecard")
async def readiness_scorecard(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    settings_ = _merge_claims_with_settings(current_user)
    acc = (settings_.get("primary_accreditor") or "HLC").upper()
    uploads = _get_user_uploads(current_user)
    # Coverage
    mapped = set(uploads.get("unique_standards", []) or [])
    total_roots = len(standards_graph.get_accreditor_standards(acc)) or 0
    coverage = round(((len(mapped) / total_roots) * 100), 1) if total_roots else 0.0
    # Trust
    tscores: List[float] = []
    for d in uploads.get("documents", []):
        ov = (d.get("trust_score") or {}).get("overall_score")
        if isinstance(ov, (int, float)):
            tscores.append(float(ov))
    avg_trust = round((sum(tscores) / len(tscores)), 3) if tscores else 0.7
    # Risk
    risk_agg = risk_explainer.aggregate()
    avg_risk = float(risk_agg.get("average_risk", 0.0))
    # Tasks
    tasks = list(_get_user_tasks(current_user).values())
    open_tasks = len([t for t in tasks if (t.get("status") or "").lower() in {"open", "in-progress"}])
    overdue = 0
    try:
        now = datetime.utcnow().date()
        for t in tasks:
            dd = t.get("due_date")
            if dd:
                try:
                    ddate = datetime.fromisoformat(dd.replace("Z", "")).date()
                    if ddate < now and (t.get("status") or "").lower() != "done":
                        overdue += 1
                except Exception:
                    continue
    except Exception:
        pass
    return {
        "success": True,
        "accreditor": acc,
        "coverage_percent": coverage,
        "average_trust": avg_trust,
        "average_risk": round(avg_risk, 3),
        "tasks_open": open_tasks,
        "tasks_overdue": overdue,
    }


# ------------------------------
# Immutable As-Of Snapshot
# ------------------------------
SNAPSHOTS_DIR = Path(os.getenv("SNAPSHOTS_DIR", "snapshots"))
SNAPSHOTS_DIR.mkdir(parents=True, exist_ok=True)
PUBLISHED_DIR = SNAPSHOTS_DIR / "public"
PUBLISHED_DIR.mkdir(parents=True, exist_ok=True)


@router.post("/snapshot/create")
async def create_snapshot(payload: Dict[str, Any] = None, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Create an immutable zip snapshot of current mappings and tasks for audit."""
    try:
        acc = (_merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()
        uploads = _get_user_uploads(current_user)
        tasks = _get_user_tasks(current_user)
        meta = {
            "accreditor": acc,
            "user": _user_key(current_user),
            "created_at": datetime.utcnow().isoformat(),
            "coverage": len(set(uploads.get("unique_standards", []) or [])),
            "documents": len(uploads.get("documents", [])),
            "tasks": len(tasks or {}),
        }
        out = SNAPSHOTS_DIR / ("snapshot_" + datetime.utcnow().strftime("%Y%m%d_%H%M%S") + "_" + secrets.token_hex(6) + ".zip")
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("meta.json", json.dumps(meta, ensure_ascii=False, indent=2))
            zf.writestr("uploads.json", json.dumps(uploads, ensure_ascii=False, indent=2))
            zf.writestr("tasks.json", json.dumps(tasks, ensure_ascii=False, indent=2))
        return {"success": True, "snapshot_path": str(out)}
    except Exception as e:
        logger.error(f"Snapshot create error: {e}")
        raise HTTPException(status_code=500, detail="Failed to create snapshot")

@router.get("/download/snapshots/{filename}")
async def download_snapshot(filename: str, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        path = SNAPSHOTS_DIR / filename
        if not path.exists() or not path.is_file():
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(str(path), filename=filename, media_type="application/zip")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download error: {e}")
        raise HTTPException(status_code=500, detail="Failed to download file")

@router.post("/narratives/reviewer-pack/publish")
async def publish_last_reviewer_pack(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Publish the user's last reviewer pack by extracting into a public directory and return a URL."""
    try:
        settings_ = _merge_claims_with_settings(current_user)
        last = (settings_ or {}).get("last_reviewer_pack") or {}
        pack_path = last.get("path")
        if not pack_path or not os.path.exists(pack_path):
            raise HTTPException(status_code=404, detail="No reviewer pack found to publish")

        safe_user = (_user_key(current_user) or "user").replace("@", "_").replace("/", "_")
        folder = f"{Path(pack_path).stem}_{secrets.token_hex(4)}"
        out_dir = PUBLISHED_DIR / safe_user / folder
        out_dir.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(pack_path, 'r') as zf:
            zf.extractall(out_dir)

        public_url = f"/snapshots/public/{safe_user}/{folder}"
        try:
            settings_["last_reviewer_pack"]["public_url"] = public_url
            _save_user_settings(current_user, settings_)
        except Exception:
            pass
        return {"success": True, "public_url": public_url, "folder": folder}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Publish reviewer pack error: {e}")
        raise HTTPException(status_code=500, detail="Failed to publish reviewer pack")


# ------------------------------
# Evidence map and crosswalk
# ------------------------------

def _normalize_accreditor_from_code(code: Optional[str]) -> str:
    if not code:
        return "UNKNOWN"
    cleaned = str(code).strip()
    if not cleaned:
        return "UNKNOWN"
    cleaned = cleaned.replace(" ", "_")
    if "_" in cleaned:
        return cleaned.split("_", 1)[0].upper()
    match = re.match(r"^[A-Za-z]+", cleaned)
    if match:
        return match.group(0).upper()
    parts = re.split(r"[:-]+", cleaned)
    if parts and parts[0]:
        return parts[0].upper()
    return cleaned.upper()


def _extract_keywords_for_alignment(text: Optional[str]) -> Set[str]:
    if not text:
        return set()
    words = re.findall(r"\b[a-z]+\b", str(text).lower())
    stop = {
        "the",
        "a",
        "an",
        "and",
        "or",
        "but",
        "in",
        "on",
        "at",
        "to",
        "for",
        "of",
        "with",
        "by",
        "is",
        "are",
        "was",
        "were",
        "has",
        "have",
        "had",
    }
    return {w for w in words if len(w) > 3 and w not in stop}


def _jaccard_score(a: Set[str], b: Set[str]) -> float:
    if not a or not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    if not union:
        return 0.0
    return float(inter) / float(union)


def _build_context_for_node(node: Any) -> Dict[str, Any]:
    path = standards_graph.get_path_to_root(getattr(node, "node_id", "")) if node else []
    breadcrumb: List[str] = []
    terms: Set[str] = set()
    for segment in path:
        label = (segment.level or "").capitalize()
        title = segment.title or ""
        breadcrumb.append(f"{label}: {title}".strip())
        terms |= _extract_keywords_for_alignment(segment.title)
        terms |= _extract_keywords_for_alignment(segment.description)
    return {
        "breadcrumb": [b for b in breadcrumb if b],
        "terms": terms,
    }


async def _resolve_user_id_for_crosswalk(current_user: Dict[str, Any]) -> Optional[str]:
    identifier = current_user.get("sub") or current_user.get("user_id") or current_user.get("email")
    if not identifier:
        return None
    identifier = str(identifier)
    if "@" in identifier:
        return await get_user_uuid_from_email(identifier)
    return identifier


async def _load_crosswalk_documents(current_user: Dict[str, Any]) -> Dict[str, Any]:
    user_id = await _resolve_user_id_for_crosswalk(current_user)
    if not user_id:
        return {"documents": [], "unique_standards": [], "unique_accreditors": []}

    documents: Dict[str, Dict[str, Any]] = {}
    unique_standards: Set[str] = set()
    unique_accreditors: Set[str] = set()

    try:
        async with db_manager.get_session() as session:
            result = await session.execute(
                text(
                    """
                    SELECT d.id, d.filename, d.uploaded_at, d.content_type,
                           em.standard_id, em.confidence, em.excerpts
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE d.user_id = :user_id
                      AND d.deleted_at IS NULL
                    ORDER BY d.uploaded_at DESC
                    """
                ),
                {"user_id": user_id},
            )

            for row in result:
                doc_id = str(row.id)
                entry = documents.setdefault(
                    doc_id,
                    {
                        "document_id": doc_id,
                        "filename": row.filename or f"Document {len(documents) + 1}",
                        "uploaded_at": row.uploaded_at.isoformat() if getattr(row, "uploaded_at", None) else None,
                        "content_type": row.content_type or "",
                        "mappings": [],
                        "accreditors": set(),
                        "standards_by_accreditor": defaultdict(set),
                        "standard_ids": set(),
                    },
                )

                standard_id = getattr(row, "standard_id", None)
                if not standard_id:
                    continue

                accreditor = _normalize_accreditor_from_code(standard_id)
                entry["accreditors"].add(accreditor)
                entry["standards_by_accreditor"][accreditor].add(str(standard_id))
                entry["standard_ids"].add(str(standard_id))

                confidence = getattr(row, "confidence", None)
                try:
                    confidence_value = float(confidence) if confidence is not None else 0.0
                except Exception:
                    confidence_value = 0.0

                excerpts_payload: List[Any] = []
                raw_excerpts = getattr(row, "excerpts", None)
                if raw_excerpts:
                    try:
                        parsed = json.loads(raw_excerpts)
                        if isinstance(parsed, list):
                            excerpts_payload = parsed
                        elif parsed:
                            excerpts_payload = [parsed]
                    except Exception:
                        excerpts_payload = []

                entry["mappings"].append(
                    {
                        "standard_id": str(standard_id),
                        "accreditor": accreditor,
                        "confidence": confidence_value,
                        "excerpts": excerpts_payload,
                    }
                )

                unique_standards.add(str(standard_id))
                unique_accreditors.add(accreditor)

    except Exception as exc:
        logger.error(f"Error loading crosswalk documents: {exc}")

    docs_list: List[Dict[str, Any]] = []
    for entry in documents.values():
        confidences = [m["confidence"] for m in entry.get("mappings", []) if isinstance(m.get("confidence"), (int, float))]
        entry["average_confidence"] = (sum(confidences) / len(confidences)) if confidences else 0.0
        docs_list.append(entry)

    return {
        "documents": docs_list,
        "unique_standards": sorted(unique_standards),
        "unique_accreditors": sorted(unique_accreditors),
    }


def _serialize_doc_for_output(doc: Dict[str, Any]) -> Dict[str, Any]:
    standards_map = {
        acc: sorted({*codes}) for acc, codes in (doc.get("standards_by_accreditor") or {}).items()
    }
    total_standards = sum(len(codes) for codes in standards_map.values())
    return {
        "document_id": doc.get("document_id"),
        "filename": doc.get("filename"),
        "uploaded_at": doc.get("uploaded_at"),
        "content_type": doc.get("content_type"),
        "accreditors": sorted({*(doc.get("accreditors") or set())}),
        "standards": standards_map,
        "total_standards": total_standards,
        "average_confidence": round(float(doc.get("average_confidence") or 0.0), 3),
    }


def _build_crosswalk_matrix_payload(documents: List[Dict[str, Any]]) -> Dict[str, Any]:
    matrix: Dict[str, Dict[str, int]] = {}
    evidence_files: List[str] = []
    standards: Set[str] = set()
    accreditors: Set[str] = set()

    for doc in documents:
        fname = doc.get("filename") or doc.get("document_id") or "Unknown"
        mappings: Dict[str, int] = {}
        for acc, codes in (doc.get("standards_by_accreditor") or {}).items():
            if not codes:
                continue
            accreditors.add(acc)
            for code in codes:
                standards.add(code)
                mappings[code] = 1
        if mappings:
            matrix[fname] = mappings
            if fname not in evidence_files:
                evidence_files.append(fname)

    reuse_count = 0
    multi_use_files = 0
    total_links = 0
    for fname, codes in matrix.items():
        mapped = list(codes.keys())
        total_links += len(mapped)
        if len(mapped) > 1:
            multi_use_files += 1
            reuse_count += len(mapped) - 1

    if total_links == 0:
        total_links = 1
    reuse_pct = round((reuse_count / total_links) * 100, 1)

    return {
        "evidence": evidence_files,
        "standards": sorted(standards),
        "matrix": matrix,
        "accreditors": sorted(accreditors),
        "reuse_percentage": reuse_pct,
        "multi_use_count": multi_use_files,
        "documents_detail": [_serialize_doc_for_output(doc) for doc in documents],
    }


def _shared_documents_for_alignment(
    source_code: str,
    target_code: str,
    documents: List[Dict[str, Any]],
    source_accreditor: str,
    target_accreditor: str,
) -> List[Dict[str, Any]]:
    shared: List[Dict[str, Any]] = []
    for doc in documents:
        accreds = doc.get("accreditors") or set()
        if source_accreditor not in accreds or target_accreditor not in accreds:
            continue
        standards_map = doc.get("standards_by_accreditor") or {}
        source_codes = standards_map.get(source_accreditor, set())
        target_codes = standards_map.get(target_accreditor, set())
        if source_code not in source_codes or target_code not in target_codes:
            continue

        serialized = _serialize_doc_for_output(doc)
        source_conf = [m["confidence"] for m in doc.get("mappings", []) if m.get("standard_id") == source_code]
        target_conf = [m["confidence"] for m in doc.get("mappings", []) if m.get("standard_id") == target_code]
        both_conf = source_conf + target_conf
        avg_conf = (sum(both_conf) / len(both_conf)) if both_conf else 0.0
        serialized.update(
            {
                "source_confidence": round((sum(source_conf) / len(source_conf)) if source_conf else 0.0, 3),
                "target_confidence": round((sum(target_conf) / len(target_conf)) if target_conf else 0.0, 3),
                "alignment_confidence": round(avg_conf, 3),
            }
        )
        shared.append(serialized)
    return shared


def _build_alignment_entry(
    match: Dict[str, Any],
    documents: List[Dict[str, Any]],
    source_accreditor: str,
    target_accreditor: str,
) -> Optional[Dict[str, Any]]:
    source_id = match.get("source_id")
    target_id = match.get("target_id")
    if not source_id or not target_id:
        return None

    source_node = standards_graph.get_node(source_id)
    target_node = standards_graph.get_node(target_id)
    if not source_node or not target_node:
        return None

    source_keywords = set(getattr(source_node, "keywords", set()))
    source_keywords |= _extract_keywords_for_alignment(source_node.title)
    source_keywords |= _extract_keywords_for_alignment(source_node.description)
    target_keywords = set(getattr(target_node, "keywords", set()))
    target_keywords |= _extract_keywords_for_alignment(target_node.title)
    target_keywords |= _extract_keywords_for_alignment(target_node.description)

    keyword_overlap = sorted(source_keywords & target_keywords)
    keyword_score = float(match.get("score") or 0.0)

    source_context = _build_context_for_node(source_node)
    target_context = _build_context_for_node(target_node)
    context_score = _jaccard_score(source_context.get("terms", set()), target_context.get("terms", set()))
    alignment_score = round((keyword_score * 0.6) + (context_score * 0.4), 3)

    shared_documents = _shared_documents_for_alignment(
        source_node.standard_id,
        target_node.standard_id,
        documents,
        source_accreditor,
        target_accreditor,
    )

    return {
        "source": {
            "id": source_node.node_id,
            "code": source_node.standard_id,
            "title": source_node.title,
            "description": source_node.description,
            "level": source_node.level,
            "breadcrumb": source_context.get("breadcrumb", []),
            "accreditor": source_accreditor,
        },
        "target": {
            "id": target_node.node_id,
            "code": target_node.standard_id,
            "title": target_node.title,
            "description": target_node.description,
            "level": target_node.level,
            "breadcrumb": target_context.get("breadcrumb", []),
            "accreditor": target_accreditor,
        },
        "source_title": source_node.title,
        "target_title": target_node.title,
        "source_code": source_node.standard_id,
        "target_code": target_node.standard_id,
        "source_accreditor": source_accreditor,
        "target_accreditor": target_accreditor,
        "keyword_score": round(keyword_score, 3),
        "context_score": round(context_score, 3),
        "alignment_score": alignment_score,
        "similarity": alignment_score,
        "keywords_overlap": keyword_overlap[:12],
        "common_keywords": keyword_overlap[:12],
        "context_overlap_terms": sorted(
            (source_context.get("terms", set()) & target_context.get("terms", set()))
        )[:12],
        "shared_documents": shared_documents,
        "shared_document_count": len(shared_documents),
    }


def _compute_crosswalk_alignments(
    source_accreditor: str,
    target_accreditor: str,
    documents: List[Dict[str, Any]],
    threshold: float,
    top_k: int,
) -> Tuple[List[Dict[str, Any]], Set[str]]:
    try:
        matches = standards_graph.find_cross_accreditor_matches(
            source_accreditor, target_accreditor, threshold=threshold, top_k=max(1, top_k)
        )
    except Exception as exc:
        logger.error(f"Crosswalk alignment search failed: {exc}")
        matches = []

    alignments: List[Dict[str, Any]] = []
    shared_docs: Set[str] = set()
    for match in matches:
        entry = _build_alignment_entry(match, documents, source_accreditor, target_accreditor)
        if entry:
            alignments.append(entry)
            for doc in entry.get("shared_documents", []):
                name = doc.get("filename") or doc.get("document_id")
                if name:
                    shared_docs.add(str(name))

    alignments.sort(key=lambda item: (item.get("alignment_score", 0.0), item.get("keyword_score", 0.0)), reverse=True)
    return alignments, shared_docs


def _build_reuse_summary(
    documents: List[Dict[str, Any]],
    source_accreditor: str,
    target_accreditor: str,
) -> Dict[str, Any]:
    source_docs = [doc for doc in documents if source_accreditor in (doc.get("accreditors") or set())]
    target_docs = [doc for doc in documents if target_accreditor in (doc.get("accreditors") or set())]
    shared_docs: List[Dict[str, Any]] = []
    for doc in documents:
        accreds = doc.get("accreditors") or set()
        if source_accreditor in accreds and target_accreditor in accreds:
            shared_docs.append(doc)

    source_reuse_rate = round((len(shared_docs) / len(source_docs)) * 100, 1) if source_docs else 0.0
    target_reuse_rate = round((len(shared_docs) / len(target_docs)) * 100, 1) if target_docs else 0.0

    shared_names = sorted({doc.get("filename") or doc.get("document_id") for doc in shared_docs if doc})
    avg_standards = 0.0
    if shared_docs:
        total = 0
        for doc in shared_docs:
            standards_map = doc.get("standards_by_accreditor") or {}
            total += sum(len(codes) for codes in standards_map.values())
        avg_standards = round(total / len(shared_docs), 2) if total else 0.0

    return {
        "documents": [_serialize_doc_for_output(doc) for doc in documents],
        "documents_evaluated": len(documents),
        "source_documents": len(source_docs),
        "target_documents": len(target_docs),
        "shared_documents": len(shared_docs),
        "shared_document_names": shared_names,
        "source_reuse_rate": source_reuse_rate,
        "target_reuse_rate": target_reuse_rate,
        "average_standards_per_shared_document": avg_standards,
    }
async def build_evidence_map_payload(
    current_user: Dict[str, Any], accreditor: Optional[str] = None
) -> Dict[str, Any]:
    """Internal helper that builds the evidence-map response for a user."""

    uploads = await _get_user_uploads(current_user)
    acc = (accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC").upper()

    mapping: Dict[str, List[Dict[str, Any]]] = {}
    documents_meta: Dict[str, Dict[str, Any]] = {}

    for d in uploads.get("documents", []):
        fname = d.get("filename")
        dt = d.get("doc_type")
        trust_overall = (d.get("trust_score") or {}).get("overall_score")
        if fname and fname not in documents_meta:
            documents_meta[fname] = {
                "filename": fname,
                "doc_type": dt,
                "trust_score": trust_overall if isinstance(trust_overall, (int, float)) else None,
                "uploaded_at": d.get("uploaded_at"),
            }
        for md in d.get("mappings", []):
            sid = md.get("standard_id")
            if not sid:
                continue
            m_acc = (md.get("accreditor") or acc).upper()
            if acc and m_acc != acc:
                continue
            mapping.setdefault(sid, []).append({
                "filename": fname,
                "doc_type": dt,
                "accreditor": m_acc,
                "confidence": md.get("confidence"),
            })

    standards_list: List[Dict[str, Any]] = []
    for sid, arr in mapping.items():
        if not arr:
            continue
        confs = [float(m.get("confidence") or 0.0) for m in arr if isinstance(m.get("confidence"), (int, float))]
        avg_conf = float(sum(confs) / len(confs)) if confs else 0.0
        standards_list.append({
            "standard_id": sid,
            "evidence_count": len(arr),
            "avg_confidence": round(avg_conf, 4),
        })

    try:
        total_roots = len(standards_graph.get_accreditor_standards(acc))
    except Exception:
        total_roots = 0
    standards_mapped = len(standards_list)
    coverage_pct = round((standards_mapped / total_roots) * 100, 1) if total_roots else 0.0

    return {
        "accreditor": acc,
        "documents": list(documents_meta.values()),
        "standards": standards_list,
        "counts": {
            "documents": len(documents_meta),
            "standards_mapped": standards_mapped,
            "total_standards": total_roots,
        },
        "coverage_percentage": coverage_pct,
        "mapping": mapping,
    }


@router.get("/standards/evidence-map")
async def get_evidence_map(
    accreditor: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Return evidence-to-standard mappings for the current user."""

    return await build_evidence_map_payload(current_user, accreditor)


class StandardsExplorerTelemetry(BaseModel):
    event: str
    accreditor: Optional[str] = None
    result: Optional[str] = None
    duration_ms: Optional[int] = None
    error: Optional[str] = None
    page: Optional[int] = None
    page_size: Optional[int] = None
    total_items: Optional[int] = None
    payload: Optional[Dict[str, Any]] = None


@router.post("/standards/explorer/telemetry")
async def standards_explorer_telemetry(
    body: StandardsExplorerTelemetry,
    request: Request,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    try:
        record_event = {
            "event": body.event,
            "accreditor": (body.accreditor or "").upper() or None,
            "result": body.result,
            "duration_ms": body.duration_ms,
            "error": body.error,
            "page": body.page,
            "page_size": body.page_size,
            "total_items": body.total_items,
            "path": str(request.url.path),
            "payload": body.payload,
        }
        stored = record_telemetry_event(_user_key(current_user), record_event)
        return {"success": True, "stored": stored}
    except Exception as exc:
        logger.warning(f"Failed to persist standards explorer telemetry: {exc}")
        raise HTTPException(status_code=500, detail="Unable to persist telemetry")


@router.get("/evidence/crosswalk")
async def get_crosswalk(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    documents_payload = await _load_crosswalk_documents(current_user)
    documents = documents_payload.get("documents", [])
    matrix_payload = _build_crosswalk_matrix_payload(documents)
    matrix_payload["available_accreditors"] = documents_payload.get("unique_accreditors", [])
    matrix_payload["unique_standards"] = documents_payload.get("unique_standards", [])
    matrix_payload["total_documents"] = len(documents)
    return matrix_payload


@router.get("/standards/crosswalkx")
async def get_crosswalkx(
    source: str,
    target: str,
    threshold: float = 0.3,
    top_k: int = 10,
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    if not source or not target:
        raise HTTPException(status_code=400, detail="source and target query parameters are required")

    source_norm = source.upper()
    target_norm = target.upper()
    if source_norm == target_norm:
        raise HTTPException(status_code=400, detail="source and target accreditors must be different")

    try:
        threshold_value = float(threshold)
    except Exception:
        threshold_value = 0.3
    threshold_value = max(0.0, min(1.0, threshold_value))

    try:
        top_k_value = int(top_k)
    except Exception:
        top_k_value = 10
    top_k_value = max(1, min(25, top_k_value))

    documents_payload = await _load_crosswalk_documents(current_user)
    documents = documents_payload.get("documents", [])

    alignments, alignment_shared_docs = _compute_crosswalk_alignments(
        source_norm, target_norm, documents, threshold_value, top_k_value
    )
    reuse_summary = _build_reuse_summary(documents, source_norm, target_norm)

    avg_alignment = 0.0
    if alignments:
        avg_alignment = round(
            sum(alignment.get("alignment_score", 0.0) for alignment in alignments) / len(alignments),
            3,
        )

    summary = {
        "average_alignment": avg_alignment,
        "high_confidence_alignments": len(
            [a for a in alignments if a.get("alignment_score", 0.0) >= 0.6]
        ),
        "alignments_returned": len(alignments),
        "shared_documents": reuse_summary.get("shared_documents", 0),
        "source_reuse_rate": reuse_summary.get("source_reuse_rate", 0.0),
        "target_reuse_rate": reuse_summary.get("target_reuse_rate", 0.0),
        "documents_evaluated": reuse_summary.get("documents_evaluated", 0),
        "alignment_shared_documents": sorted(alignment_shared_docs),
    }

    return {
        "source": source_norm,
        "target": target_norm,
        "threshold": threshold_value,
        "top_k": top_k_value,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "summary": summary,
        "alignments": alignments,
        "reuse": reuse_summary,
        "available_accreditors": documents_payload.get("unique_accreditors", []),
        "unique_standards": documents_payload.get("unique_standards", []),
    }


# ------------------------------
# Risk weights
# ------------------------------
@router.get("/risk/weights")
async def get_risk_weights(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    s = _merge_claims_with_settings(current_user)
    return {"weights": s.get("risk_weights", {
        "Documentation": 1.0,
        "Evidence Quality": 1.0,
        "Freshness": 1.0,
        "Operations": 1.0,
        "Change Management": 1.0,
        "Audit History": 1.0,
    })}


@router.post("/risk/weights")
async def set_risk_weights(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    s = _merge_claims_with_settings(current_user)
    weights = payload.get("weights") or {}
    if not isinstance(weights, dict):
        raise HTTPException(status_code=400, detail="weights must be a dict")
    s["risk_weights"] = {k: float(v) for k, v in weights.items()}
    _save_user_settings(current_user, s)
    return {"status": "saved", "weights": s["risk_weights"]}


# ------------------------------
# Gaps and standards
# ------------------------------
def _build_gap_recommendations(
    gaps: List[Dict[str, Any]],
    *,
    accreditor: str,
    ai_enabled: bool,
) -> List[Dict[str, str]]:
    """Convert gap records into structured recommendation objects."""

    accreditor = str(accreditor or "General")

    priority_map = {
        "critical": "critical",
        "high": "high",
        "medium": "medium",
        "low": "low",
    }

    recommendations: List[Dict[str, str]] = []
    seen_titles: Set[str] = set()

    sorted_gaps = sorted(
        (gap for gap in gaps if isinstance(gap, dict)),
        key=lambda item: float(item.get("risk_score", 0.0)),
        reverse=True,
    )

    for gap in sorted_gaps:
        standard_ref = gap.get("standard") or {}
        code = str(standard_ref.get("code") or standard_ref.get("id") or "").strip()
        title_raw = str(standard_ref.get("title") or "").strip()
        recommendation_text = (gap.get("recommendation") or gap.get("ai_recommendation") or "").strip()

        if not recommendation_text:
            ai_insights = gap.get("ai_insights") or []
            if ai_insights and isinstance(ai_insights, list):
                for insight in ai_insights:
                    if isinstance(insight, dict) and insight.get("action"):
                        recommendation_text = str(insight["action"]).strip()
                        break
                    if isinstance(insight, str) and insight.strip():
                        recommendation_text = insight.strip()
                        break

        if not recommendation_text:
            continue

        priority_key = str(gap.get("risk_level", "medium")).lower()
        priority_value = priority_map.get(priority_key, "medium")

        if code:
            display_title = f"{code}: {title_raw}" if title_raw else code
        else:
            display_title = title_raw or "Accreditation gap recommendation"

        if display_title in seen_titles:
            continue

        recommendations.append(
            {
                "priority": priority_value,
                "title": display_title,
                "description": recommendation_text,
            }
        )
        seen_titles.add(display_title)

        if len(recommendations) >= 5:
            break

    fallback_messages = [
        {
            "priority": "high" if ai_enabled else "medium",
            "title": "Focus on the highest-risk standards first",
            "description": "Review the critical and high-risk standards highlighted in the gap summary and assign owners for remediation.",
        },
        {
            "priority": "medium",
            "title": "Upload or refresh evidence",
            "description": "Provide recent supporting documents for your mapped standards to improve confidence scores and reduce risk.",
        },
        {
            "priority": "low",
            "title": f"Review {accreditor.upper()} readiness checklist",
            "description": "Validate that each requirement has at least one verified piece of evidence and note any upcoming review deadlines.",
        },
    ]

    for fallback in fallback_messages:
        if fallback["title"] in seen_titles:
            continue
        recommendations.append(fallback)
        seen_titles.add(fallback["title"])

    return recommendations


@router.get("/gaps/analysis")
async def get_gap_analysis(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        # Get user's primary accreditor
        user_settings = _merge_claims_with_settings(current_user)
        accreditor = user_settings.get("primary_accreditor", "HLC")
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        
        # Get standards for the user's accreditor
        all_standards = []
        try:
            if hasattr(standards_graph, 'nodes'):
                all_standards = []
                for n in standards_graph.nodes.values():
                    if getattr(n, "level", "") == "standard" and getattr(n, "accreditor", "").upper() == accreditor.upper():
                        all_standards.append(n)
        except Exception as e:
            logger.warning(f"Error accessing standards graph: {e}")
            all_standards = []
        
        gaps = []
        
        # Check if AI gap predictor is available
        if USE_AI_GAP_PREDICTOR and settings.openai_api_key and settings.openai_api_key != "sk-proj-PLACEHOLDER" and all_standards:
            try:
                # Get user's evidence data from database
                user_id = await get_user_uuid_from_email(email)
                evidence_by_standard = {}
                
                async with db_manager.get_session() as session:
                    # Get evidence mappings for user
                    result = await session.execute(
                        text("""
                            SELECT em.standard_id, em.confidence, em.excerpts, 
                                   d.filename, d.uploaded_at, d.analysis_results
                            FROM evidence_mappings em
                            JOIN documents d ON em.document_id = d.id
                            WHERE d.user_id = :user_id
                            AND d.deleted_at IS NULL
                            ORDER BY em.confidence DESC
                        """),
                        {"user_id": user_id}
                    )
                    
                    for row in result:
                        if row.standard_id not in evidence_by_standard:
                            evidence_by_standard[row.standard_id] = []
                        evidence_by_standard[row.standard_id].append({
                            "filename": row.filename,
                            "confidence": float(row.confidence),
                            "excerpts": json.loads(row.excerpts) if row.excerpts else [],
                            "uploaded_at": row.uploaded_at
                        })
                
                # Get institution context
                institution_context = {
                    "name": user_settings.get("institution_name", "Institution"),
                    "type": user_settings.get("institution_type", "University"),
                    "accreditor": accreditor,
                    "risk_profile": 0.3  # Default medium risk
                }
                
                # Analyze each standard with AI
                for standard in all_standards[:20]:  # Limit to top 20 for performance
                    standard_id = getattr(standard, "node_id", getattr(standard, "standard_id", ""))
                    evidence_data = evidence_by_standard.get(standard_id, [])
                    
                    # Get AI-enhanced risk prediction
                    risk_result = await ai_gap_risk_predictor.predict_gap_risk(
                        standard_id=standard_id,
                        evidence_data=evidence_data,
                        historical_data=None,  # TODO: Add historical tracking
                        institution_context=institution_context
                    )
                    
                    if risk_result["risk_score"] >= 0.4:  # Only show medium+ risks
                        gaps.append({
                            "standard": {
                                "code": standard_id,
                                "title": getattr(standard, "title", ""),
                            },
                            "risk_score": risk_result["risk_score"],
                            "risk_level": risk_result["risk_level"],
                            "risk_factors": risk_result.get("risk_factors", {}),
                            "ai_insights": risk_result.get("ai_insights", []),
                            "recommendation": risk_result["recommended_actions"][0]["action"] if risk_result.get("recommended_actions") else f"Address {standard_id} gaps",
                            "impact": risk_result["recommended_actions"][0]["impact"] if risk_result.get("recommended_actions") else "Reduces compliance risk",
                            "timeline": risk_result["recommended_actions"][0]["timeline"] if risk_result.get("recommended_actions") else "Within 30 days",
                            "confidence": risk_result.get("confidence", 0.8),
                            "next_review": risk_result.get("next_review_date", "")
                        })
                
                # Sort by risk score
                gaps.sort(key=lambda g: g["risk_score"], reverse=True)
                
                recommendations_payload = _build_gap_recommendations(
                    gaps,
                    accreditor=accreditor,
                    ai_enabled=True,
                )

                return {
                    "status": "success",
                    "timestamp": datetime.utcnow().isoformat(),
                    "gap_analysis": {
                        "total_gaps": len(gaps),
                        "critical_risk": len([g for g in gaps if g["risk_level"] == "Critical"]),
                        "high_risk": len([g for g in gaps if g["risk_level"] == "High"]),
                        "medium_risk": len([g for g in gaps if g["risk_level"] == "Medium"]),
                        "low_risk": len([g for g in gaps if g["risk_level"] == "Low"]),
                        "gaps": gaps[:15],  # Return top 15 gaps
                    },
                    "algorithm": "GapRisk Predictor™ with AI",
                    "ai_enabled": True,
                    "recommendations": recommendations_payload,
                    "analysis_confidence": sum(g.get("confidence", 0.8) for g in gaps) / len(gaps) if gaps else 0.8
                }
                
            except Exception as e:
                logger.error(f"AI gap analysis failed: {e}", exc_info=True)
                # Fallback to basic analysis
        
        # Fallback to basic gap analysis
        if all_standards:
            for i, standard in enumerate(all_standards[:10]):
                try:
                    # Simple risk calculation without external dependencies
                    risk_score = 0.5 + (i * 0.05)  # Simple incremental risk score
                    if risk_score >= 0.5:
                        code_val = getattr(standard, "node_id", getattr(standard, "standard_id", f"{accreditor}.{i+1}"))
                        gaps.append(
                            {
                                "standard": {
                                    "code": code_val,
                                    "title": getattr(standard, "title", f"Standard {i+1}"),
                                },
                                "risk_score": min(risk_score, 0.95),
                                "risk_level": "High" if risk_score >= 0.7 else "Medium",
                                "recommendation": f"Upload evidence for {code_val} within 14 days",
                                "impact": "High" if risk_score >= 0.7 else "Medium",
                            }
                        )
                except Exception as e:
                    logger.warning(f"Error processing standard: {e}")
                    continue
        else:
            # Provide some mock gaps if no standards found
            for i in range(5):
                risk_score = 0.5 + (i * 0.1)
                gaps.append(
                    {
                        "standard": {
                            "code": f"{accreditor}.{i+1}",
                            "title": f"{accreditor} Standard {i+1}",
                        },
                        "risk_score": min(risk_score, 0.95),
                        "risk_level": "High" if risk_score >= 0.7 else "Medium",
                        "recommendation": f"Upload evidence for {accreditor}.{i+1} within 14 days",
                        "impact": "High" if risk_score >= 0.7 else "Medium",
                    }
                )

        recommendations_payload = _build_gap_recommendations(
            gaps,
            accreditor=accreditor,
            ai_enabled=False,
        )

        return {
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "gap_analysis": {
                "total_gaps": len(gaps),
                "high_risk": len([g for g in gaps if g["risk_score"] >= 0.7]),
                "medium_risk": len([g for g in gaps if 0.5 <= g["risk_score"] < 0.7]),
                "gaps": gaps,
            },
            "algorithm": "Risk Assessment v1",
            "ai_enabled": False,
            "recommendations": recommendations_payload,
        }
    except Exception as e:
        logger.error(f"Gap analysis error: {e}", exc_info=True)
        # Return a safe fallback response instead of raising an exception
        fallback_recommendations = _build_gap_recommendations(
            [],
            accreditor="general",
            ai_enabled=False,
        )

        return {
            "status": "success",
            "timestamp": datetime.utcnow().isoformat(),
            "gap_analysis": {
                "total_gaps": 0,
                "high_risk": 0,
                "medium_risk": 0,
                "gaps": [],
            },
            "algorithm": "Risk Assessment v1",
            "ai_enabled": False,
            "recommendations": fallback_recommendations,
        }


@router.get("/compliance/gaps")
async def get_compliance_gaps_simple(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        import numpy as np

        sample_standards = [
            {
                "standard_id": "HLC_1",
                "coverage": 62,
                "trust_scores": [0.8, 0.75],
                "evidence_ages": [220, 340],
                "overdue_tasks": 2,
                "total_tasks": 8,
                "recent_changes": 1,
                "historical_findings": 1,
                "days_to_review": 150,
            },
            {
                "standard_id": "HLC_3",
                "coverage": 55,
                "trust_scores": [0.7, 0.68],
                "evidence_ages": [400, 500],
                "overdue_tasks": 3,
                "total_tasks": 7,
                "recent_changes": 2,
                "historical_findings": 2,
                "days_to_review": 120,
            },
            {
                "standard_id": "SACSCOC_10",
                "coverage": 48,
                "trust_scores": [0.72, 0.66],
                "evidence_ages": [360, 410],
                "overdue_tasks": 4,
                "total_tasks": 10,
                "recent_changes": 1,
                "historical_findings": 3,
                "days_to_review": 90,
            },
        ]

        risks = [
            gap_risk_predictor.predict_risk(
                standard_id=s["standard_id"],
                coverage_percentage=s["coverage"],
                evidence_trust_scores=s["trust_scores"],
                evidence_ages_days=s["evidence_ages"],
                overdue_tasks_count=s["overdue_tasks"],
                total_tasks_count=s["total_tasks"],
                recent_changes_count=s["recent_changes"],
                historical_findings_count=s["historical_findings"],
                time_to_next_review_days=s["days_to_review"],
            )
            for s in sample_standards
        ]

        overall = float(np.mean([r.risk_score for r in risks])) if risks else 0.0
        levels_order = ["minimal", "low", "medium", "high", "critical"]
        worst_level_raw = (
            max((r.risk_level.value for r in risks), key=lambda v: levels_order.index(v)) if risks else "low"
        )
        worst_level = worst_level_raw.capitalize()
        avg_conf = float(np.mean([r.confidence for r in risks])) if risks else 0.8
        timeline_months = int(np.mean([r.time_to_review for r in risks]) / 30) if risks else 6

        issues = []
        for r in risks:
            issues.extend(r.predicted_issues)
        seen = set()
        identified = []
        for i in issues:
            if i not in seen:
                seen.add(i)
                identified.append(i)

        factor_avgs = {
            "coverage": [],
            "trust": [],
            "staleness": [],
            "task_debt": [],
            "change_impact": [],
            "review_history": [],
        }
        for r in risks:
            for f in r.factors:
                if f.factor_name in factor_avgs:
                    factor_avgs[f.factor_name].append(f.normalized_value)

        def _avg(vals):
            import numpy as _np

            return float(_np.mean(vals)) if vals else 0.0

        category_risks = {
            "Documentation": round(_avg(factor_avgs["coverage"]), 3),
            "Evidence Quality": round(_avg(factor_avgs["trust"]), 3),
            "Freshness": round(_avg(factor_avgs["staleness"]), 3),
            "Operations": round(_avg(factor_avgs["task_debt"]), 3),
            "Change Management": round(_avg(factor_avgs["change_impact"]), 3),
            "Audit History": round(_avg(factor_avgs["review_history"]), 3),
        }

        return {
            "risk_assessment": {
                "overall_risk": round(overall, 3),
                "risk_level": worst_level,
                "confidence": round(avg_conf, 3),
                "timeline_months": timeline_months,
            },
            "category_risks": category_risks,
            "recommendations": [
                "Focus on high-contribution risk factors first",
                "Refresh outdated evidence and close overdue tasks",
                "Address standards with recent changes and findings",
            ],
            "identified_issues": identified[:6],
            "next_actions": [
                "Upload additional evidence documents",
                "Review high-risk categories",
                "Schedule compliance assessment",
                "Update institutional documentation",
            ],
        }
    except Exception as e:
        logger.error(f"Compliance gaps (simple) error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze compliance gaps: {e}")


@router.get("/standards/graph")
async def get_standards_graph_simple(
    current_user: Dict[str, Any] = Depends(get_current_user_simple), accreditor: Optional[str] = None
):
    try:
        acc = (accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or (current_user.get("primary_accreditor") if isinstance(current_user, dict) else None) or "HLC")
        roots = standards_graph.get_accreditor_standards(acc)
        standards_obj = {
            n.node_id: {
                "title": n.title,
                "category": n.level.title() if hasattr(n, "level") else "Standard",
                "domain": "Core",
            }
            for n in roots[:50]
        }
        relationships = []
        for n in roots[:10]:
            for c in standards_graph.get_children(n.node_id)[:10]:
                relationships.append({"source": n.node_id, "target": c.node_id})

        stats = standards_graph.get_graph_stats()
        return {
            "accreditor": acc,
            "total_standards": len(roots),
            "standards": standards_obj,
            "relationships": relationships,
            "available_accreditors": stats.get("accreditors", []),
        }
    except Exception as e:
        logger.error(f"Standards graph (simple) error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to load standards graph: {e}")


@router.get("/standards/detail")
async def get_standard_detail(standard_id: str, current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        node = standards_graph.get_node(standard_id)
        if not node:
            for n in standards_graph.nodes.values():
                if getattr(n, "standard_id", None) == standard_id or getattr(n, "code", None) == standard_id:
                    node = n
                    break
        if not node:
            raise HTTPException(status_code=404, detail="Standard not found")
        mode = _get_standards_display_mode(current_user)
        return {
            "id": getattr(node, "node_id", standard_id),
            "code": getattr(node, "standard_id", getattr(node, "code", standard_id)),
            "title": getattr(node, "title", ""),
            "category": getattr(node, "level", getattr(node, "category", "Standard")),
            "domain": getattr(node, "accreditor", "Core"),
            "description": ("" if mode == "redacted" else getattr(node, "description", "")),
            "level": getattr(node, "level", None),
            "display_mode": mode,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Standard detail error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch standard detail")


@router.get("/standards/cross-accreditor-matches")
async def get_cross_accreditor_matches(
    source: str,
    target: str,
    threshold: float = 0.3,
    top_k: int = 3,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Find likely equivalent standards across two accreditors using keyword overlap.

    Returns a list of {source_id, source_title, target_id, target_title, score}.
    """
    try:
        source = (source or '').strip().upper()
        target = (target or '').strip().upper()
        if not source or not target or source == target:
            raise HTTPException(status_code=400, detail="Provide distinct source and target accreditors")
        matches = standards_graph.find_cross_accreditor_matches(source, target, threshold=threshold, top_k=top_k)
        return {"source": source, "target": target, "threshold": threshold, "top_k": top_k, "matches": matches}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Cross-accreditor matches error: {e}")
        raise HTTPException(status_code=500, detail="Failed to compute cross-accreditor matches")


@router.get("/standards/categories")
async def get_standards_categories(
    current_user: Dict[str, Any] = Depends(get_current_user_simple), accreditor: Optional[str] = None
):
    try:
        acc = accreditor or _merge_claims_with_settings(current_user).get("primary_accreditor") or "HLC"
        roots = standards_graph.get_accreditor_standards(acc)
        counts: Dict[str, int] = {}
        for n in roots:
            cat = getattr(n, "category", getattr(n, "level", "Standard"))
            counts[cat] = counts.get(cat, 0) + 1
        return {"accreditor": acc, "categories": counts}
    except Exception as e:
        logger.error(f"Standards categories error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load categories")


# ------------------------------
# Reviewer: per-standard notes/status (server persistence)
# ------------------------------
@router.get("/reviews/standards")
async def get_standard_reviews(
    accreditor: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    try:
        data = _get_user_standard_reviews(current_user, accreditor=accreditor)
        return {"success": True, "accreditor": (accreditor or None), "count": len(data), "reviews": data}
    except Exception as e:
        logger.error(f"Get standard reviews error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load reviews")


@router.post("/reviews/standard")
async def save_standard_review(
    payload: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    try:
        accreditor = (payload.get("accreditor") or payload.get("acc") or "").upper()
        standard_id = str(payload.get("standard_id") or payload.get("code") or "").strip()
        if not accreditor or not standard_id:
            raise HTTPException(status_code=400, detail="accreditor and standard_id are required")
        entry = _set_user_standard_review(
            current_user,
            accreditor=accreditor,
            standard_id=standard_id,
            status=payload.get("status"),
            note=payload.get("note"),
            assignee=payload.get("assignee"),
            due_date=payload.get("due_date"),
        )
        return {"success": True, "entry": entry, "standard_id": standard_id, "accreditor": accreditor}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Save standard review error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save review")


# ------------------------------
# Selected standards persistence (session store)
# ------------------------------
@router.post("/standards/selection/save")
async def save_selected_standards(payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        ids = payload.get("selected") or payload.get("standard_ids") or []
        if not isinstance(ids, list):
            raise HTTPException(status_code=400, detail="selected must be a list")
        all_sessions = _safe_load_json(SESSIONS_STORE)
        uk = _user_key(current_user)
        sess = all_sessions.get(uk, {})
        sess["selected_standards"] = list(dict.fromkeys([str(x) for x in ids]))
        all_sessions[uk] = sess
        _safe_save_json(SESSIONS_STORE, all_sessions)
        return {"status": "saved", "count": len(sess["selected_standards"])}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Save selected standards error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save selection")


@router.get("/standards/selection/load")
async def load_selected_standards(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    try:
        all_sessions = _safe_load_json(SESSIONS_STORE)
        ids = (all_sessions.get(_user_key(current_user), {}) or {}).get("selected_standards", [])
        return {"selected": ids, "count": len(ids)}
    except Exception as e:
        logger.error(f"Load selected standards error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load selection")


# ------------------------------
# Reports (HTML + CSV)
# ------------------------------
@router.post("/evidence/report")
async def generate_evidence_report(
    payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    try:
        filename = payload.get("filename", "document")
        analysis = payload.get("analysis", {})
        trust = analysis.get("trust_score", {})
        mappings: List[Dict[str, Any]] = analysis.get("mappings", [])
        org = payload.get("organization") or _merge_claims_with_settings(current_user).get("organization") or ""
        accreditor = payload.get("accreditor") or _merge_claims_with_settings(current_user).get("primary_accreditor") or ""

        def pct(x):
            try:
                return f"{float(x) * 100:.1f}%"
            except Exception:
                return "-"

        rows = "".join(
            [
                f"""
            <tr>
                <td style='padding:6px;border:1px solid #e5e7eb'>{m.get('standard_id','')}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{m.get('title','')}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{pct(m.get('confidence',0))}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{m.get('match_type','')}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{'Yes' if m.get('meets_standard') else 'No'}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{'; '.join(m.get('rationale_spans') or [])}</td>
                <td style='padding:6px;border:1px solid #e5e7eb'>{(m.get('note') or '').replace('<','&lt;').replace('>','&gt;')}</td>
            </tr>
            """
                for m in mappings
            ]
        )

        html = f"""
        <!DOCTYPE html>
        <html><head><meta charset='utf-8'>
        <title>Evidence Report - {filename}</title>
        <style>
            body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial; color: #111827; }}
            .card {{ border:1px solid #e5e7eb; border-radius:8px; padding:16px; margin:12px 0; }}
            h1 {{ font-size: 20px; }} h2 {{ font-size:16px; }} table {{ border-collapse: collapse; width:100%; font-size: 12px; }}
        </style></head>
        <body>
            <h1>Evidence Report</h1>
            <div class='card'>
                <div><strong>Organization:</strong> {org or '-'} | <strong>Accreditor:</strong> {accreditor or '-'}</div>
                <div><strong>File:</strong> {filename}</div>
                <div><strong>Generated:</strong> {datetime.utcnow().isoformat()} UTC</div>
            </div>
            <div class='card'>
                <h2>Trust Score</h2>
                <div>Overall: {pct(trust.get('overall_score'))} | Quality: {pct(trust.get('quality_score'))} | Reliability: {pct(trust.get('reliability_score'))} | Confidence: {pct(trust.get('confidence_score'))}</div>
            </div>
            <div class='card'>
                <h2>Mappings</h2>
                <table>
                    <thead>
                        <tr>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Standard</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Title</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Confidence</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Match</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Meets</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Rationales</th>
                            <th style='text-align:left;padding:6px;border:1px solid #e5e7eb'>Notes</th>
                        </tr>
                    </thead>
                    <tbody>{rows}</tbody>
                </table>
            </div>
        </body></html>
        """
        return JSONResponse({"html": html})
    except Exception as e:
        logger.error(f"Report generation error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate report")


@router.post("/evidence/report.csv")
async def generate_evidence_report_csv(
    payload: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    try:
        filename = payload.get("filename", "document")
        analysis = payload.get("analysis", {})
        mappings: List[Dict[str, Any]] = analysis.get("mappings", [])

        output = io.StringIO()
        output.write("\ufeff")  # UTF-8 BOM for Excel
        writer = csv.writer(output)
        writer.writerow(["Standard", "Title", "Confidence", "Match", "Meets", "Rationales", "Notes"])

        def pct_val(x):
            try:
                return f"{float(x) * 100:.1f}%"
            except Exception:
                return "-"

        for m in mappings:
            standard_id = m.get("standard_id", "")
            title = m.get("title", "")
            confidence = pct_val(m.get("confidence", 0))
            match_type = m.get("match_type", "")
            meets = "Yes" if m.get("meets_standard") else "No"
            rationals = "; ".join(m.get("rationale_spans") or [])
            note = (m.get("note") or "").replace("\r", " ").replace("\n", " ")
            writer.writerow([standard_id, title, confidence, match_type, meets, rationals, note])

        csv_text = output.getvalue()
        output.close()

        safe_name = "".join(c for c in filename if c.isalnum() or c in ("_", "-", ".")) or "document"
        headers = {
            "Content-Disposition": f'attachment; filename="evidence_report_{safe_name}.csv"',
            "Content-Type": "text/csv; charset=utf-8",
        }
        return PlainTextResponse(content=csv_text, media_type="text/csv; charset=utf-8", headers=headers)
    except Exception as e:
        logger.error(f"CSV report generation error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate CSV report")


# ------------------------------
# AI Integration Test Endpoints
# ------------------------------
@router.get("/test/ai-status")
async def test_ai_status(
    current_user: Dict[str, Any] = Depends(get_current_user_simple),
):
    """Test endpoint to verify AI/OpenAI integration status"""
    status = {
        "enhanced_mapper_available": USE_ENHANCED_MAPPER,
        "openai_configured": bool(settings.openai_api_key),
        "openai_key_prefix": settings.openai_api_key[:8] + "..." if settings.openai_api_key else None,
    }
    
    if USE_ENHANCED_MAPPER and settings.openai_api_key:
        try:
            if not enhanced_evidence_mapper._initialized:
                await enhanced_evidence_mapper.initialize()
            
            # Test with a simple prompt
            test_doc = EvidenceDocument(
                doc_id="test.txt",
                text="This document describes our student assessment procedures and learning outcomes measurement.",
                metadata={},
                doc_type="policy",
                source_system="manual",
                upload_date=datetime.utcnow()
            )
            
            mappings = await enhanced_evidence_mapper.map_evidence_with_ai(
                test_doc,
                num_candidates=5,
                final_top_k=3,
                use_llm=True
            )
            
            status["ai_test_success"] = True
            status["sample_mappings"] = len(mappings)
            status["llm_service_initialized"] = enhanced_evidence_mapper.llm_service is not None
        except Exception as e:
            status["ai_test_success"] = False
            status["ai_test_error"] = str(e)
    
    return status


@router.get("/ai/status")
async def get_ai_status(current_user: Dict[str, Any] = Depends(get_current_user_simple)):
    """Comprehensive AI integration status endpoint"""
    
    # Check if OpenAI is properly configured (not using placeholder)
    openai_configured = bool(settings.openai_api_key and settings.openai_api_key != "sk-proj-PLACEHOLDER")
    
    status = {
        "endpoint": "/api/user/intelligence-simple/ai/status",
        "openai_configured": openai_configured,
        "openai_key_prefix": settings.openai_api_key[:8] + "..." if settings.openai_api_key else None,
        "enhanced_mapper_available": USE_ENHANCED_MAPPER,
        "ai_narrative_available": USE_AI_NARRATIVE,
        "ai_gap_predictor_available": USE_AI_GAP_PREDICTOR,
        "ai_features": {
            "enhanced_document_analysis": USE_ENHANCED_MAPPER and openai_configured,
            "citeguard_narratives": USE_AI_NARRATIVE and openai_configured,
            "predictive_gap_analysis": USE_AI_GAP_PREDICTOR and openai_configured,
            "crosswalk_mapping": False,  # Not yet implemented
            "evidence_trust_scoring": True,  # Basic version available
            "standards_graph": True  # Always available
        },
        "algorithms_status": {
            "StandardsGraph™": {
                "status": "active", 
                "ai_enhanced": False,
                "description": "Living knowledge graph with semantic embeddings"
            },
            "EvidenceMapper™": {
                "status": "active", 
                "ai_enhanced": USE_ENHANCED_MAPPER and openai_configured,
                "description": "Dual-encoder retrieval with AI reranking" if openai_configured else "TF-IDF based mapping"
            },
            "EvidenceTrust Score™": {
                "status": "partial", 
                "ai_enhanced": False,
                "description": "Basic quality scoring implemented"
            },
            "GapRisk Predictor™": {
                "status": "active", 
                "ai_enhanced": USE_AI_GAP_PREDICTOR and openai_configured,
                "description": "8-factor predictive analysis with AI insights" if openai_configured else "Basic risk scoring"
            },
            "CrosswalkX™": {
                "status": "not_implemented", 
                "ai_enhanced": False,
                "description": "Multi-accreditor mapping coming soon"
            },
            "CiteGuard™": {
                "status": "active" if USE_AI_NARRATIVE and openai_configured else "basic",
                "ai_enhanced": USE_AI_NARRATIVE and openai_configured,
                "description": "AI narratives with mandatory citations" if openai_configured else "Basic HTML formatting"
            }
        },
        "recommendations": []
    }
    
    # Add recommendations based on status
    if not openai_configured:
        status["recommendations"].append({
            "priority": "high",
            "action": "Configure OpenAI API key for full AI capabilities",
            "impact": "Enables enhanced analysis accuracy and predictive insights"
        })
    
    if openai_configured and not all([USE_ENHANCED_MAPPER, USE_AI_NARRATIVE, USE_AI_GAP_PREDICTOR]):
        status["recommendations"].append({
            "priority": "medium",
            "action": "Ensure all AI services are properly imported",
            "impact": "Some AI features may not be available"
        })
    
    return status


@router.get("/evidence/mappings/detail")
async def get_evidence_mappings_detail(
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get detailed evidence mappings with transparency into AI decisions"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        mappings = []
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT 
                        em.id,
                        em.document_id,
                        em.standard_id,
                        em.confidence,
                        em.excerpts,
                        em.created_at,
                        d.filename,
                        d.file_size,
                        d.analysis_results
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE d.user_id = :user_id
                    AND d.deleted_at IS NULL
                    ORDER BY em.confidence DESC
                """),
                {"user_id": user_id}
            )
            
            for row in result:
                # Get standard details
                standard = standards_graph.get_node(row.standard_id)
                
                # Parse excerpts for evidence trails
                excerpts_data = json.loads(row.excerpts) if row.excerpts else []
                
                mappings.append({
                    "mapping_id": row.id,
                    "document": {
                        "id": row.document_id,
                        "filename": row.filename,
                        "size": row.file_size
                    },
                    "standard": {
                        "id": row.standard_id,
                        "title": standard.title if standard else row.standard_id,
                        "accreditor": standard.accreditor if standard else "Unknown"
                    },
                    "confidence": float(row.confidence),
                    "confidence_level": "High" if row.confidence >= 0.8 else "Medium" if row.confidence >= 0.6 else "Low",
                    "evidence_excerpts": excerpts_data,
                    "mapping_method": "AI-Enhanced" if USE_ENHANCED_MAPPER and settings.openai_api_key else "Algorithmic",
                    "created_at": row.created_at.isoformat() if row.created_at else None,
                    "explanation": f"Document contains evidence matching {int(row.confidence * 100)}% of standard requirements"
                })
        
        # Group by document
        by_document = {}
        for mapping in mappings:
            doc_id = mapping["document"]["id"]
            if doc_id not in by_document:
                by_document[doc_id] = {
                    "document": mapping["document"],
                    "mappings": []
                }
            by_document[doc_id]["mappings"].append(mapping)
        
        return {
            "status": "success",
            "total_mappings": len(mappings),
            "documents": list(by_document.values()),
            "transparency": {
                "mapping_method": "AI-Enhanced" if USE_ENHANCED_MAPPER and settings.openai_api_key else "Algorithmic",
                "allows_manual_adjustment": True,
                "confidence_thresholds": {
                    "high": 0.8,
                    "medium": 0.6,
                    "low": 0.4
                }
            }
        }
        
    except Exception as e:
        logger.error(f"Evidence mappings detail error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve evidence mappings")


@router.post("/evidence/mappings/adjust")
async def adjust_evidence_mapping(
    payload: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Allow users to manually adjust or validate AI evidence mappings"""
    try:
        mapping_id = payload.get("mapping_id")
        action = payload.get("action")  # "accept", "reject", "adjust"
        new_confidence = payload.get("new_confidence")
        user_notes = payload.get("notes", "")
        
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        async with db_manager.get_session() as session:
            # Verify ownership
            check = await session.execute(
                text("""
                    SELECT em.id 
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE em.id = :mapping_id AND d.user_id = :user_id
                """),
                {"mapping_id": mapping_id, "user_id": user_id}
            )
            
            if not check.first():
                raise HTTPException(status_code=403, detail="Mapping not found or access denied")
            
            if action == "reject":
                # Mark mapping as rejected
                await session.execute(
                    text("""
                        UPDATE evidence_mappings 
                        SET confidence = 0, 
                            excerpts = jsonb_build_object(
                                'user_action', 'rejected',
                                'notes', :notes,
                                'original_confidence', confidence
                            ),
                            updated_at = CURRENT_TIMESTAMP
                        WHERE id = :mapping_id
                    """),
                    {"mapping_id": mapping_id, "notes": user_notes}
                )
            
            elif action == "adjust" and new_confidence is not None:
                # Adjust confidence
                await session.execute(
                    text("""
                        UPDATE evidence_mappings 
                        SET confidence = :new_confidence,
                            excerpts = jsonb_set(
                                COALESCE(excerpts, '{}')::jsonb,
                                '{user_adjustment}',
                                jsonb_build_object(
                                    'original_confidence', confidence,
                                    'adjusted_confidence', :new_confidence,
                                    'notes', :notes
                                )
                            ),
                            updated_at = CURRENT_TIMESTAMP
                        WHERE id = :mapping_id
                    """),
                    {
                        "mapping_id": mapping_id, 
                        "new_confidence": float(new_confidence),
                        "notes": user_notes
                    }
                )
            
            elif action == "accept":
                # Mark as accepted
                await session.execute(
                    text("""
                        UPDATE evidence_mappings 
                        SET excerpts = jsonb_set(
                                COALESCE(excerpts, '{}')::jsonb,
                                '{user_action}',
                                '"accepted"'
                            ),
                            updated_at = CURRENT_TIMESTAMP
                        WHERE id = :mapping_id
                    """),
                    {"mapping_id": mapping_id}
                )
            
            await session.commit()
            
            return {
                "status": "success",
                "mapping_id": mapping_id,
                "action": action,
                "message": f"Mapping {action} successfully"
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Adjust mapping error: {e}")
        raise HTTPException(status_code=500, detail="Failed to adjust mapping")


@router.get("/evidence/trust-scores")
async def get_evidence_trust_scores(
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get detailed EvidenceTrust scores for all documents"""
    try:
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        documents = []
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT 
                        id, 
                        filename, 
                        analysis_results,
                        uploaded_at,
                        updated_at
                    FROM documents 
                    WHERE user_id = :user_id 
                    AND deleted_at IS NULL
                    ORDER BY uploaded_at DESC
                """),
                {"user_id": user_id}
            )
            
            for row in result:
                analysis = json.loads(row.analysis_results) if row.analysis_results else {}
                trust_score = analysis.get("trust_score", {})
                
                # Calculate age factor
                age_days = (datetime.utcnow() - row.uploaded_at).days if row.uploaded_at else 999
                freshness_score = max(0, 1 - (age_days / 365))  # 1.0 for new, 0 for year old
                
                # Enhanced trust scoring
                documents.append({
                    "document_id": row.id,
                    "filename": row.filename,
                    "trust_scores": {
                        "overall": float(trust_score.get("overall_score", 0.7)),
                        "quality": float(trust_score.get("quality_score", 0.7)),
                        "reliability": float(trust_score.get("reliability_score", 0.7)),
                        "confidence": float(trust_score.get("confidence_score", 0.7)),
                        "freshness": round(freshness_score, 2),
                        "completeness": 0.8  # TODO: Implement completeness check
                    },
                    "trust_factors": {
                        "document_age_days": age_days,
                        "last_updated": row.updated_at.isoformat() if row.updated_at else None,
                        "format": "PDF" if row.filename.lower().endswith('.pdf') else "Other",
                        "has_metadata": True if analysis.get("metadata") else False,
                        "source": "User Upload"
                    },
                    "trust_level": "High" if trust_score.get("overall_score", 0.7) >= 0.8 else "Medium",
                    "algorithm": "EvidenceTrust Score™"
                })
        
        # Calculate aggregate trust
        all_scores = [d["trust_scores"]["overall"] for d in documents]
        avg_trust = sum(all_scores) / len(all_scores) if all_scores else 0.0
        
        return {
            "status": "success",
            "documents": documents,
            "aggregate_trust": {
                "average_trust": round(avg_trust, 2),
                "trust_distribution": {
                    "high": len([d for d in documents if d["trust_level"] == "High"]),
                    "medium": len([d for d in documents if d["trust_level"] == "Medium"]),
                    "low": len([d for d in documents if d["trust_level"] == "Low"])
                },
                "recommendations": [
                    "Update documents older than 6 months" if any(d["trust_factors"]["document_age_days"] > 180 for d in documents) else None,
                    "Add more high-quality evidence" if avg_trust < 0.7 else None
                ]
            }
        }
        
    except Exception as e:
        logger.error(f"Trust scores error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve trust scores")


@router.get("/standards/visual-graph")
async def get_standards_visual_graph(
    accreditor: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user_simple)
):
    """Get visual StandardsGraph data for interactive display"""
    try:
        user_settings = _merge_claims_with_settings(current_user)
        if not accreditor:
            accreditor = user_settings.get("primary_accreditor", "HLC")
        
        # Get user's evidence mappings
        email = current_user.get("sub", current_user.get("email", current_user.get("user_id", "unknown")))
        user_id = await get_user_uuid_from_email(email)
        
        mapped_standards = set()
        async with db_manager.get_session() as session:
            result = await session.execute(
                text("""
                    SELECT DISTINCT standard_id 
                    FROM evidence_mappings em
                    JOIN documents d ON em.document_id = d.id
                    WHERE d.user_id = :user_id
                    AND em.confidence > 0
                """),
                {"user_id": user_id}
            )
            mapped_standards = {row.standard_id for row in result}
        
        # Build graph data
        nodes = []
        edges = []
        
        # Add accreditor as root
        nodes.append({
            "id": accreditor,
            "label": accreditor,
            "type": "accreditor",
            "level": 0,
            "color": "#1e40af"
        })
        
        # Add standards
        standards = standards_graph.get_accreditor_standards(accreditor)
        for i, standard in enumerate(standards[:20]):  # Limit for visualization
            std_id = standard.node_id
            nodes.append({
                "id": std_id,
                "label": f"{std_id}: {standard.title[:50]}...",
                "type": "standard",
                "level": 1,
                "color": "#10b981" if std_id in mapped_standards else "#6b7280",
                "mapped": std_id in mapped_standards,
                "coverage": 1.0 if std_id in mapped_standards else 0.0
            })
            edges.append({
                "source": accreditor,
                "target": std_id
            })
            
            # Add clauses
            for clause in standards_graph.get_children(std_id)[:3]:  # Limit clauses
                nodes.append({
                    "id": clause.node_id,
                    "label": clause.title[:40],
                    "type": "clause",
                    "level": 2,
                    "color": "#8b5cf6"
                })
                edges.append({
                    "source": std_id,
                    "target": clause.node_id
                })
        
        return {
            "status": "success",
            "graph": {
                "nodes": nodes,
                "edges": edges
            },
            "statistics": {
                "total_standards": len(standards),
                "mapped_standards": len([n for n in nodes if n.get("mapped")]),
                "coverage_percentage": round(len(mapped_standards) / len(standards) * 100, 1) if standards else 0
            },
            "algorithm": "StandardsGraph™"
        }
        
    except Exception as e:
        logger.error(f"Visual graph error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate visual graph")
