"""
Document Processing Service for A3E

Handles file upload, text extraction, and document processing for evidence items.
Supports PDF, DOCX, XLSX, CSV, TXT, and MD files.
"""

import asyncio
import logging
import tempfile
import os
from typing import Dict, Any, Optional, List
from pathlib import Path
import mimetypes
from datetime import datetime
import uuid

# Initialize logger early
logger = logging.getLogger(__name__)

# Optional document processing libraries
try:
    import pypdf
    import docx
    import pandas as pd
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False
    logger.warning("Document processing libraries not available - limited file support")
    
    # Mock pandas for basic functionality
    class pd:
        @staticmethod
        def DataFrame(data):
            return {"data": data}
        
        @staticmethod 
        def read_excel(file):
            return {"data": "mock"}

from fastapi import UploadFile

from ..core.config import Settings
from ..models import Evidence, EvidenceType, ProcessingStatus
from ..services.database_service import DatabaseService


class DocumentService:
    """Document processing service for evidence files"""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'application/vnd.ms-excel': self._process_xlsx,
            'text/csv': self._process_csv,
            'text/plain': self._process_text,
            'text/markdown': self._process_text,
        }
    
    async def process_uploaded_file(
        self,
        file: UploadFile,
        institution_id: str,
        evidence_type: str,
        description: Optional[str] = None,
        uploaded_by: str = "system"
    ) -> Evidence:
        """Process an uploaded file and create evidence record"""
        
        # Validate file type
        mime_type = file.content_type or mimetypes.guess_type(file.filename)[0]
        if mime_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        # Create evidence record
        evidence_data = {
            "id": str(uuid.uuid4()),
            "institution_id": institution_id,
            "title": file.filename or "Uploaded Document",
            "description": description,
            "evidence_type": EvidenceType(evidence_type),
            "file_name": file.filename,
            "file_size_bytes": file.size,
            "mime_type": mime_type,
            "processing_status": ProcessingStatus.PENDING,
            "created_at": datetime.utcnow()
        }
        
        # Save initial record to database
        db_service = DatabaseService(self.settings.database_url)
        await db_service.initialize()
        
        try:
            evidence = await db_service.create_evidence(evidence_data)
            
            # Process file in background
            asyncio.create_task(
                self._process_file_background(evidence.id, file, mime_type, db_service)
            )
            
            return evidence
            
        finally:
            await db_service.close()
    
    async def _process_file_background(
        self,
        evidence_id: str,
        file: UploadFile,
        mime_type: str,
        db_service: DatabaseService
    ):
        """Process file content in background"""
        try:
            # Update status to processing
            await db_service.update_evidence_status(
                evidence_id,
                ProcessingStatus.PROCESSING
            )
            
            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text based on file type
                processor = self.supported_types.get(mime_type)
                if processor:
                    extracted_data = await processor(temp_file_path)
                else:
                    raise ValueError(f"No processor for mime type: {mime_type}")
                
                # Update evidence with extracted content
                await self._update_evidence_with_content(
                    evidence_id,
                    extracted_data,
                    db_service
                )
                
                # Update status to completed
                await db_service.update_evidence_status(
                    evidence_id,
                    ProcessingStatus.COMPLETED
                )
                
                logger.info(f"Successfully processed evidence {evidence_id}")
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to process evidence {evidence_id}: {e}")
            await db_service.update_evidence_status(
                evidence_id,
                ProcessingStatus.FAILED,
                str(e)
            )
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            text_content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                        "creator": pdf_reader.metadata.get("/Creator", ""),
                        "producer": pdf_reader.metadata.get("/Producer", ""),
                        "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")),
                        "modification_date": str(pdf_reader.metadata.get("/ModDate", ""))
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            return {
                "extracted_text": text_content.strip(),
                "structured_data": {
                    "metadata": metadata,
                    "page_count": len(pdf_reader.pages),
                    "file_type": "pdf"
                },
                "keywords": self._extract_keywords(text_content)
            }
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            table_data = []
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                
                if table_text.strip():
                    table_data.append(table_text)
                    text_content += "\n--- Table ---\n" + table_text + "\n"
            
            # Extract core properties
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "last_modified_by": core_props.last_modified_by or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            return {
                "extracted_text": text_content.strip(),
                "structured_data": {
                    "metadata": metadata,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "tables": table_data,
                    "file_type": "docx"
                },
                "keywords": self._extract_keywords(text_content)
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Extract data from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            text_content = ""
            structured_data = {
                "sheets": {},
                "file_type": "xlsx"
            }
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to text representation
                    sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
                    sheet_text += df.to_string(index=False) + "\n"
                    text_content += sheet_text
                    
                    # Store structured data
                    structured_data["sheets"][sheet_name] = {
                        "shape": df.shape,
                        "columns": list(df.columns),
                        "data_types": df.dtypes.to_dict(),
                        "sample_data": df.head(5).to_dict() if not df.empty else {},
                        "summary_stats": df.describe().to_dict() if not df.empty else {}
                    }
                    
                except Exception as e:
                    logger.warning(f"Failed to process sheet {sheet_name}: {e}")
                    continue
            
            return {
                "extracted_text": text_content.strip(),
                "structured_data": structured_data,
                "keywords": self._extract_keywords(text_content)
            }
            
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            raise
    
    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Could not decode CSV file with any supported encoding")
            
            # Convert to text representation
            text_content = df.to_string(index=False)
            
            # Structured data
            structured_data = {
                "shape": df.shape,
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "sample_data": df.head(10).to_dict() if not df.empty else {},
                "summary_stats": df.describe().to_dict() if not df.empty else {},
                "file_type": "csv"
            }
            
            return {
                "extracted_text": text_content,
                "structured_data": structured_data,
                "keywords": self._extract_keywords(text_content)
            }
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            raise
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text or markdown file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise ValueError("Could not decode text file with any supported encoding")
            
            # Basic statistics
            lines = text_content.split('\n')
            words = text_content.split()
            
            structured_data = {
                "line_count": len(lines),
                "word_count": len(words),
                "character_count": len(text_content),
                "file_type": "text"
            }
            
            return {
                "extracted_text": text_content,
                "structured_data": structured_data,
                "keywords": self._extract_keywords(text_content)
            }
            
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise
    
    def _extract_keywords(self, text: str, max_keywords: int = 20) -> List[str]:
        """Extract keywords from text using simple frequency analysis"""
        if not text:
            return []
        
        try:
            # Simple keyword extraction (can be enhanced with NLP libraries)
            words = text.lower().split()
            
            # Filter common stop words
            stop_words = {
                'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
                'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
                'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must',
                'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they',
                'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'its', 'our', 'their'
            }
            
            # Clean and filter words
            cleaned_words = []
            for word in words:
                # Remove punctuation and check length
                word = ''.join(c for c in word if c.isalnum()).strip()
                if len(word) > 3 and word not in stop_words:
                    cleaned_words.append(word)
            
            # Count frequency
            word_freq = {}
            for word in cleaned_words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Sort by frequency and return top keywords
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            return [word for word, freq in sorted_words[:max_keywords] if freq > 1]
            
        except Exception as e:
            logger.error(f"Keyword extraction failed: {e}")
            return []
    
    async def _update_evidence_with_content(
        self,
        evidence_id: str,
        extracted_data: Dict[str, Any],
        db_service: DatabaseService
    ):
        """Update evidence record with extracted content"""
        try:
            update_data = {
                "extracted_text": extracted_data.get("extracted_text"),
                "structured_data": extracted_data.get("structured_data"),
                "keywords": extracted_data.get("keywords", []),
                "relevance_score": 0.8,  # Default relevance score
                "confidence_score": 0.9,  # Default confidence score
                "processed_at": datetime.utcnow()
            }
            
            # Remove None values
            update_data = {k: v for k, v in update_data.items() if v is not None}
            
            await db_service.execute(
                """
                UPDATE evidence 
                SET extracted_text = :extracted_text,
                    structured_data = :structured_data,
                    keywords = :keywords,
                    relevance_score = :relevance_score,
                    confidence_score = :confidence_score,
                    processed_at = :processed_at,
                    updated_at = :updated_at
                WHERE id = :evidence_id
                """,
                {
                    **update_data,
                    "evidence_id": evidence_id,
                    "updated_at": datetime.utcnow()
                }
            )
            
        except Exception as e:
            logger.error(f"Failed to update evidence content: {e}")
            raise
    
    async def get_supported_file_types(self) -> Dict[str, List[str]]:
        """Get list of supported file types and extensions"""
        return {
            "mime_types": list(self.supported_types.keys()),
            "extensions": [
                ".pdf", ".docx", ".xlsx", ".xls", 
                ".csv", ".txt", ".md"
            ],
            "descriptions": {
                "pdf": "Portable Document Format",
                "docx": "Microsoft Word Document",
                "xlsx": "Microsoft Excel Spreadsheet",
                "csv": "Comma-Separated Values",
                "txt": "Plain Text",
                "md": "Markdown Document"
            }
        }
